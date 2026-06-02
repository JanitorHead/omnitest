import streamlit as st
import requests
import re
import xml.etree.ElementTree as ET
import zipfile
import tempfile
import os
import base64
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import genanki
from fpdf import FPDF

_GEMINI_OK = True  # Solo necesita 'requests', ya instalado

st.set_page_config(
    page_title="Daypo Extractor",
    page_icon="📝",
    layout="centered",
    initial_sidebar_state="collapsed"
)


# ---------------------------------------------------------------------------
# Funciones auxiliares
# ---------------------------------------------------------------------------

def limpiar_nombre_carpeta(texto: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", texto).strip()


def nombre_archivo(texto: str) -> str:
    return re.sub(r'\s+', '_', limpiar_nombre_carpeta(texto)).strip('_')


def obtener_imagen(id_test: str, num_imagen: str, url_origen: str) -> bytes | None:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Referer": url_origen,
        "Accept": "image/webp,image/apng,image/*,*/*;q=0.8",
    }
    prefijo = id_test[:3]
    url = f"https://www.daypo.com/testimages/{prefijo}/{id_test}_{num_imagen}.jpg"
    try:
        r = requests.get(url, headers=headers, timeout=8)
        if r.status_code == 200:
            return r.content
    except requests.RequestException:
        pass
    return None


def procesar_pregunta(doc: Document, num: int, enunciado: str,
                      img_bytes: bytes | None, opciones: list[tuple[str, bool]],
                      con_respuesta: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {enunciado}")
    run.bold = True
    run.font.size = Pt(11)

    if img_bytes:
        doc.add_picture(BytesIO(img_bytes), width=Inches(3.5))

    for texto_op, es_correcta in opciones:
        p_op = doc.add_paragraph()
        if con_respuesta and es_correcta:
            r = p_op.add_run(f"   - {texto_op}  (correcta)")
            r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x73, 0x28)
        else:
            p_op.add_run(f"   - {texto_op}")

    doc.add_paragraph()


def extraer_datos_test(url: str, headers_web: dict) -> tuple[str, str, list] | None:
    try:
        res_html = requests.get(url, headers=headers_web, timeout=10)
        res_html.raise_for_status()
        match = re.search(r"ntest\s*=\s*(\d+)", res_html.text)
        if not match:
            return None
        id_test = match.group(1)
    except requests.RequestException:
        return None

    try:
        res_xml = requests.post(
            "https://www.daypo.com/asps/load.php",
            data={"tes": id_test},
            headers=headers_web,
            timeout=10,
        )
        res_xml.raise_for_status()
        raiz = ET.fromstring(res_xml.text)
    except (requests.RequestException, ET.ParseError):
        return None

    nodo_titulo = raiz.find(".//t")
    titulo = nodo_titulo.text.strip() if (nodo_titulo is not None and nodo_titulo.text) else f"Test {id_test}"

    contenedor = raiz.find("c")
    if contenedor is None:
        return None

    preguntas = []
    for q in contenedor.findall("c"):
        nodo_p = q.find("p")
        enunciado = nodo_p.text.strip() if nodo_p is not None and nodo_p.text else "Sin enunciado"
        nodo_b = q.find("b")
        num_imagen = nodo_b.text.strip() if nodo_b is not None and nodo_b.text else None
        nodo_c = q.find("c")
        mascara = nodo_c.text if nodo_c is not None and nodo_c.text else ""
        indice_correcta = mascara.find("2")
        nodo_r = q.find("r")
        opciones = []
        if nodo_r is not None:
            for idx, nodo_o in enumerate(nodo_r.findall("o")):
                texto = nodo_o.text.strip() if nodo_o.text else ""
                opciones.append((texto, idx == indice_correcta))
        preguntas.append({"enunciado": enunciado, "num_imagen": num_imagen, "opciones": opciones})

    return id_test, titulo, preguntas


def extraer_enlaces_daypo(texto: str) -> list[str]:
    patron = r'https?://(?:www\.)?daypo\.com/[^\s\n\r"\'<>()\[\]{}]+'
    candidatos = [re.sub(r'[.,;:!?]+$', '', e) for e in re.findall(patron, texto)]
    vistos: set[str] = set()
    resultado = []
    for e in candidatos:
        if e not in vistos:
            vistos.add(e)
            resultado.append(e)
    return resultado


def generar_nombre_base(tests_datos: list[dict]) -> str:
    if not tests_datos:
        return "Daypo"
    if len(tests_datos) == 1:
        return nombre_archivo(tests_datos[0]["titulo"])
    titulos = [t["titulo"] for t in tests_datos]
    conteo: dict[str, int] = {}
    for titulo in titulos:
        for palabra in set(re.findall(r'[a-zA-ZÀ-ɏ]{4,}', titulo.lower())):
            conteo[palabra] = conteo.get(palabra, 0) + 1
    umbral = max(2, len(titulos) // 2)
    candidatos = [(p, c) for p, c in conteo.items() if c >= umbral]
    if candidatos:
        palabra = max(candidatos, key=lambda x: x[1])[0]
        return f"{palabra.capitalize()}_{len(titulos)}_tests"
    return f"{nombre_archivo(titulos[0])[:25].rstrip('_')}_y_{len(titulos) - 1}_mas"


# ---------------------------------------------------------------------------
# Exportadores Word
# ---------------------------------------------------------------------------

def generar_word_combinado(todos_los_tests: list[dict], con_respuesta: bool = True) -> bytes:
    """Genera un .docx con todos los tests. Para 1 test devuelve el ya generado si con_respuesta=True."""
    if len(todos_los_tests) == 1 and con_respuesta:
        return todos_los_tests[0]["_word_bytes"]

    doc = Document()
    if len(todos_los_tests) > 1:
        doc.add_heading("Recopilacion completa de tests de Daypo", 0)

    for test in todos_los_tests:
        doc.add_heading(test["titulo"], 1)
        for i, p in enumerate(test["preguntas"], 1):
            procesar_pregunta(doc, i, p["enunciado"], p.get("img_bytes"),
                              p["opciones"], con_respuesta=con_respuesta)
        if len(todos_los_tests) > 1:
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_zip_word_individuales(todos_los_tests: list[dict]) -> bytes:
    """ZIP con un .docx por test, en subcarpetas Con_Respuesta/ y Sin_Respuesta/."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for t in todos_los_tests:
            nb = nombre_archivo(t["titulo"])
            zf.writestr(f"Con_Respuesta/{nb}.docx", t["_word_bytes"])

            doc_sin = Document()
            doc_sin.add_heading(t["titulo"], 1)
            for i, p in enumerate(t["preguntas"], 1):
                procesar_pregunta(doc_sin, i, p["enunciado"], p.get("img_bytes"),
                                  p["opciones"], con_respuesta=False)
            buf_sin = BytesIO()
            doc_sin.save(buf_sin)
            zf.writestr(f"Sin_Respuesta/{nb}.docx", buf_sin.getvalue())
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Exportadores PDF
# ---------------------------------------------------------------------------

def generar_pdf(tests_datos: list[dict], con_respuesta: bool = True) -> bytes:
    def safe(t: str) -> str:
        return t.encode("latin-1", errors="replace").decode("latin-1")

    MARGEN, W = 20, 170
    pdf = FPDF(format="A4")
    pdf.set_margins(MARGEN, MARGEN, MARGEN)
    pdf.set_auto_page_break(auto=True, margin=MARGEN)

    for test in tests_datos:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_x(MARGEN)
        try:
            pdf.multi_cell(W, 10, safe(test["titulo"]), align="L")
        except Exception:
            pass
        pdf.ln(4)

        for i, p in enumerate(test["preguntas"], 1):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_x(MARGEN)
            try:
                pdf.multi_cell(W, 8, safe(f"{i}. {p['enunciado']}"), align="L")
            except Exception:
                pass

            if p.get("img_bytes"):
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
                    f.write(p["img_bytes"])
                    tmp = f.name
                try:
                    pdf.image(tmp, x=MARGEN, w=100)
                    pdf.ln(2)
                except Exception:
                    pass
                finally:
                    os.unlink(tmp)

            for texto, es_correcta in p["opciones"]:
                if not texto:
                    continue
                pdf.set_x(MARGEN + 4)
                if con_respuesta and es_correcta:
                    pdf.set_font("Helvetica", "B", 11)
                    try:
                        pdf.multi_cell(W - 4, 7, safe(f">> {texto}"), align="L")
                    except Exception:
                        pass
                else:
                    pdf.set_font("Helvetica", "", 11)
                    try:
                        pdf.multi_cell(W - 4, 7, safe(texto), align="L")
                    except Exception:
                        pass
            pdf.ln(4)

    return bytes(pdf.output())


def generar_zip_pdf_individuales(todos_los_tests: list[dict]) -> bytes:
    """ZIP con un .pdf por test, en subcarpetas Con_Respuesta/ y Sin_Respuesta/."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for t in todos_los_tests:
            nb = nombre_archivo(t["titulo"])
            zf.writestr(f"Con_Respuesta/{nb}.pdf", generar_pdf([t], con_respuesta=True))
            zf.writestr(f"Sin_Respuesta/{nb}.pdf", generar_pdf([t], con_respuesta=False))
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Exportadores RemNote e Imagenes
# ---------------------------------------------------------------------------

def generar_zip_remnote(tests_datos: list[dict]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        lineas: list[str] = []
        for test in tests_datos:
            lineas.append(test["titulo"])
            lineas.append("")
            for p in test["preguntas"]:
                correcta = None
                incorrectas: list[str] = []
                for texto, es_correcta in p["opciones"]:
                    if not texto:
                        continue
                    if es_correcta:
                        correcta = texto
                    else:
                        incorrectas.append(texto)
                if correcta is None:
                    continue
                if p.get("img_bytes") and p.get("img_nombre"):
                    nombre = p["img_nombre"]
                    zf.writestr(f"images/{nombre}", p["img_bytes"])
                    lineas.append(f"- **{p['enunciado']}")
                    lineas.append(f"**![](images/{nombre}) >>A)")
                else:
                    lineas.append(f"- **{p['enunciado']}** >>A)")
                lineas.append(f"    - {correcta}")
                for inc in incorrectas:
                    lineas.append(f"    - {inc}")
                lineas.append("")
            lineas.append("")
        zf.writestr("Banco_de_Preguntas_MCQ.md", "\n".join(lineas).encode("utf-8"))
    buf.seek(0)
    return buf.getvalue()


def generar_zip_imagenes(tests_datos: list[dict]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for test in tests_datos:
            carpeta = nombre_archivo(test["titulo"])
            n = 1
            for p in test["preguntas"]:
                if p.get("img_bytes"):
                    zf.writestr(f"{carpeta}/{carpeta}_{n:03d}.jpg", p["img_bytes"])
                    n += 1
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Anki
# ---------------------------------------------------------------------------

_ANKI_MODEL = genanki.Model(
    1607392323,
    "Daypo MCQ Interactive",
    fields=[
        {"name": "Question"}, {"name": "Image"},
        {"name": "Option A"}, {"name": "Option B"}, {"name": "Option C"},
        {"name": "Option D"}, {"name": "Option E"}, {"name": "Correct Answer"},
    ],
    templates=[{
        "name": "MCQ Card",
        "qfmt": """\
<div class="card-container">
  <div class="question">{{Question}}</div>
  {{#Image}}<div class="image">{{Image}}</div>{{/Image}}
  <hr>
  <div class="options-container" id="options-box">
    {{#Option A}}<button class="option-btn" id="btn-A" onclick="selectOption('A')">{{Option A}}</button>{{/Option A}}
    {{#Option B}}<button class="option-btn" id="btn-B" onclick="selectOption('B')">{{Option B}}</button>{{/Option B}}
    {{#Option C}}<button class="option-btn" id="btn-C" onclick="selectOption('C')">{{Option C}}</button>{{/Option C}}
    {{#Option D}}<button class="option-btn" id="btn-D" onclick="selectOption('D')">{{Option D}}</button>{{/Option D}}
    {{#Option E}}<button class="option-btn" id="btn-E" onclick="selectOption('E')">{{Option E}}</button>{{/Option E}}
  </div>
</div>
<script>
function selectOption(letter) {
  var btns = document.getElementsByClassName("option-btn");
  for (var i = 0; i < btns.length; i++) btns[i].classList.remove("selected");
  document.getElementById("btn-" + letter).classList.add("selected");
  sessionStorage.setItem("ankiUserChoice", letter);
}
(function initOptions() {
  var box = document.getElementById("options-box");
  var btns = Array.from(box.children);
  var saved = sessionStorage.getItem("ankiOrder");
  if (saved) {
    saved.split(",").forEach(function(id) { var el=document.getElementById(id); if(el)box.appendChild(el); });
  } else {
    sessionStorage.removeItem("ankiUserChoice");
    for (var i=btns.length-1;i>0;i--){var j=Math.floor(Math.random()*(i+1));var t=btns[i];btns[i]=btns[j];btns[j]=t;}
    sessionStorage.setItem("ankiOrder", btns.map(function(b){return b.id;}).join(","));
    btns.forEach(function(b){box.appendChild(b);});
  }
})();
</script>""",
        "afmt": """\
{{FrontSide}}
<div id="correct-key" style="display:none;">{{Correct Answer}}</div>
<script>
var correct=document.getElementById("correct-key").innerText.trim().toUpperCase();
var chosen=sessionStorage.getItem("ankiUserChoice")||"";
var cBtn=document.getElementById("btn-"+correct);
if(cBtn)cBtn.classList.add("correct");
if(chosen&&chosen!==correct){var wBtn=document.getElementById("btn-"+chosen);if(wBtn)wBtn.classList.add("incorrect");}
sessionStorage.removeItem("ankiOrder");
</script>""",
    }],
    css="""\
.card{font-family:Arial,Helvetica,sans-serif;font-size:16px;text-align:left;color:#1a1a1a;background-color:#f0f0f0;}
.card-container{max-width:620px;margin:20px auto;padding:20px;background:#fff;border-radius:10px;box-shadow:0 4px 10px rgba(0,0,0,.12);}
.question{font-size:1.2em;font-weight:bold;margin-bottom:14px;color:#1a1a1a;}
.image{margin:12px 0;}.image img{max-width:100%;max-height:280px;border-radius:6px;}
.options-container{display:flex;flex-direction:column;gap:10px;}
.option-btn{text-align:left;padding:13px 16px;border:2px solid #d0d0d0;border-radius:8px;background:#fafafa;color:#1a1a1a;cursor:pointer;font-size:1em;transition:all .18s ease;}
.option-btn:hover{background:#efefef;border-color:#aaa;}
.option-btn.selected{border-color:#0056b3;background:#dceeff;color:#003a80;}
.option-btn.correct{border-color:#28a745!important;background:#d4edda!important;color:#155724!important;}
.option-btn.incorrect{border-color:#dc3545!important;background:#f8d7da!important;color:#721c24!important;}
.card.nightMode{color:#e8e8e8;background-color:#1e1e1e;}
.card.nightMode .card-container{background:#2a2a2a;}
.card.nightMode .question{color:#e8e8e8;}
.card.nightMode .option-btn{background:#333;border-color:#555;color:#e8e8e8;}
.card.nightMode .option-btn:hover{background:#3d3d3d;border-color:#777;}
.card.nightMode .option-btn.selected{border-color:#5aabff;background:#1a3a5c;color:#a8d4ff;}
.card.nightMode .option-btn.correct{border-color:#2ecc71!important;background:#1a3d2b!important;color:#7fdca0!important;}
.card.nightMode .option-btn.incorrect{border-color:#e74c3c!important;background:#3d1a1a!important;color:#f5a0a0!important;}
""",
)


def generar_apkg_anki(tests_datos: list[dict]) -> bytes:
    deck = genanki.Deck(2059400110, "Daypo Extractor")
    with tempfile.TemporaryDirectory() as tmpdir:
        media_files: list[str] = []
        for test in tests_datos:
            for p in test["preguntas"]:
                correcta = None
                incorrectas: list[str] = []
                for texto, es_correcta in p["opciones"]:
                    if not texto:
                        continue
                    if es_correcta:
                        correcta = texto
                    else:
                        incorrectas.append(texto)
                if correcta is None:
                    continue
                campos = [correcta] + incorrectas
                fields_ops = [campos[i] if i < len(campos) else "" for i in range(5)]
                image_html = ""
                if p.get("img_bytes") and p.get("img_nombre"):
                    nombre = p["img_nombre"]
                    img_path = os.path.join(tmpdir, nombre)
                    with open(img_path, "wb") as f:
                        f.write(p["img_bytes"])
                    media_files.append(img_path)
                    image_html = f'<img src="{nombre}">'
                note = genanki.Note(
                    model=_ANKI_MODEL,
                    fields=[p["enunciado"], image_html, *fields_ops, "A"],
                    guid=genanki.guid_for(test["titulo"], p["enunciado"]),
                )
                deck.add_note(note)
        pkg = genanki.Package(deck)
        pkg.media_files = media_files
        apkg_path = os.path.join(tmpdir, "daypo.apkg")
        pkg.write_to_file(apkg_path)
        with open(apkg_path, "rb") as f:
            return f.read()


# ---------------------------------------------------------------------------
# HTML Quiz
# ---------------------------------------------------------------------------

_QUIZ_CSS = """
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',system-ui,sans-serif;background:#0f172a;color:#e2e8f0;min-height:100vh;padding:16px 12px}
#app{max-width:1100px;margin:0 auto}
.hdr{display:flex;justify-content:space-between;align-items:center;background:#1e293b;border-radius:12px;padding:12px 16px;margin-bottom:12px}
.hdr-title{font-size:13px;color:#94a3b8;font-weight:500;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;max-width:60%}
.hdr-score{font-size:13px;color:#60a5fa;font-weight:700;white-space:nowrap}
.pbar{width:100%;height:5px;background:#1e293b;border-radius:99px;margin-bottom:16px;overflow:hidden}
.pfill{height:100%;background:linear-gradient(90deg,#3b82f6,#8b5cf6);border-radius:99px;transition:width .3s ease}
.main{display:flex;gap:16px;align-items:flex-start}
.quiz-area{flex:1;min-width:0}
.card{background:#1e293b;border-radius:14px;padding:24px;box-shadow:0 4px 20px rgba(0,0,0,.3)}
.qnum{font-size:11px;color:#64748b;font-weight:600;text-transform:uppercase;letter-spacing:.06em;margin-bottom:10px}
.qtext{font-size:17px;font-weight:600;line-height:1.55;color:#f1f5f9;margin-bottom:16px}
.qimg{width:100%;max-height:260px;object-fit:contain;border-radius:8px;margin-bottom:16px;border:1px solid #334155}
.opts{display:flex;flex-direction:column;gap:8px;margin-bottom:16px}
.opt{padding:12px 15px;background:#0f172a;border:2px solid #334155;border-radius:9px;color:#cbd5e1;font-size:14px;text-align:left;cursor:pointer;transition:all .14s ease;width:100%}
.opt:hover:not(:disabled){border-color:#60a5fa;background:#1e3a5f;color:#e2e8f0}
.opt.sel{border-color:#3b82f6;background:#1d4ed8;color:#fff}
.opt.ok{border-color:#22c55e!important;background:#14532d!important;color:#bbf7d0!important}
.opt.bad{border-color:#ef4444!important;background:#7f1d1d!important;color:#fecaca!important}
.opt:disabled{cursor:default}
.navbar{display:flex;gap:8px;flex-wrap:wrap}
.btn{padding:10px 14px;border:none;border-radius:8px;font-size:13px;font-weight:600;cursor:pointer;transition:all .14s ease}
.btn:disabled{opacity:.4;cursor:default}
.btn-prev{background:#334155;color:#94a3b8;flex-shrink:0}.btn-prev:hover:not(:disabled){background:#475569;color:#e2e8f0}
.btn-skip{background:#334155;color:#94a3b8}.btn-skip:hover:not(:disabled){background:#475569;color:#e2e8f0}
.btn-show{background:#475569;color:#e2e8f0;flex:1}.btn-show:hover:not(:disabled){background:#64748b}
.btn-next{background:#3b82f6;color:#fff;flex-shrink:0}.btn-next:hover:not(:disabled){background:#2563eb}
.kbd{display:inline-block;background:#0f172a;color:#64748b;font-size:10px;padding:1px 4px;border-radius:3px;font-family:monospace;margin-left:3px}
.grid-panel{width:180px;flex-shrink:0;background:#1e293b;border-radius:14px;padding:14px;position:sticky;top:16px}
.gp-title{font-size:11px;color:#64748b;font-weight:600;text-transform:uppercase;letter-spacing:.06em;margin-bottom:10px}
.grid{display:grid;grid-template-columns:repeat(5,1fr);gap:5px}
.gq{aspect-ratio:1;border:none;border-radius:5px;font-size:11px;font-weight:700;cursor:pointer;transition:all .12s;color:#94a3b8;background:#0f172a;display:flex;align-items:center;justify-content:center}
.gq:hover{filter:brightness(1.3)}.gq.cur{outline:2px solid #60a5fa;outline-offset:1px;color:#e2e8f0}
.gq.ok{background:#15803d;color:#bbf7d0}.gq.bad{background:#b91c1c;color:#fecaca}
.legend{margin-top:10px;display:flex;flex-direction:column;gap:5px}
.li{display:flex;align-items:center;gap:6px;font-size:11px;color:#64748b}
.ld{width:9px;height:9px;border-radius:2px;flex-shrink:0}
.btn-change{width:100%;margin-top:12px;padding:8px;background:#0f172a;border:none;border-radius:7px;color:#64748b;font-size:11px;font-weight:600;cursor:pointer;transition:all .14s}
.btn-change:hover{background:#334155;color:#94a3b8}
.sel-screen{text-align:center;padding:32px 16px}
.sel-title{font-size:22px;font-weight:700;color:#f1f5f9;margin-bottom:6px}
.sel-sub{font-size:14px;color:#64748b;margin-bottom:24px}
.test-list{display:flex;flex-direction:column;gap:10px;margin-bottom:28px;text-align:left}
.test-item{display:flex;align-items:center;gap:12px;padding:14px 16px;background:#0f172a;border:2px solid #334155;border-radius:10px;cursor:pointer;transition:border-color .15s}
.test-item:hover{border-color:#475569}.test-item.on{border-color:#3b82f6}
.test-item input[type=checkbox]{width:17px;height:17px;cursor:pointer;accent-color:#3b82f6;flex-shrink:0}
.t-name{font-size:14px;font-weight:600;color:#e2e8f0}.t-count{font-size:12px;color:#64748b;margin-top:2px}
.btn-start{background:#3b82f6;color:#fff;border:none;padding:14px 36px;border-radius:10px;font-size:15px;font-weight:700;cursor:pointer}.btn-start:hover{background:#2563eb}
.end{text-align:center;padding:36px 16px}
.end-pct{font-size:68px;font-weight:800;background:linear-gradient(135deg,#3b82f6,#8b5cf6);-webkit-background-clip:text;-webkit-text-fill-color:transparent}
.end-sub{font-size:18px;color:#94a3b8;margin:8px 0 28px}
.end-btns{display:flex;gap:12px;justify-content:center;flex-wrap:wrap}
.btn-restart{background:#3b82f6;color:#fff;padding:12px 26px;border-radius:9px;font-size:14px;font-weight:700;cursor:pointer;border:none}.btn-restart:hover{background:#2563eb}
@media(max-width:640px){.main{flex-direction:column-reverse}.grid-panel{width:100%;position:static}.grid{grid-template-columns:repeat(8,1fr)}}
"""

_QUIZ_JS = r"""
const TESTS=QUIZ_TESTS_DATA;const TITULO="QUIZ_TITLE";const MULTI=TESTS.length>1;
let selIdx=TESTS.map((_,i)=>i),qs=[],ans=[],rev=[],cho=[],cur=0;
const app=document.getElementById('app');
function shuf(a){const b=[...a];for(let i=b.length-1;i>0;i--){const j=Math.floor(Math.random()*(i+1));[b[i],b[j]]=[b[j],b[i]];}return b;}
const KEY='dq_'+TESTS.reduce((s,t)=>s+t.nombre.slice(0,3),'')+TESTS.reduce((s,t)=>s+t.preguntas.length,0);
function save(){try{localStorage.setItem(KEY,JSON.stringify({qs,ans,rev,cho,cur,selIdx}));}catch(e){}}
function loadSaved(){try{const sv=JSON.parse(localStorage.getItem(KEY)||'null');if(sv&&sv.qs&&sv.qs.length>0){qs=sv.qs;ans=sv.ans;rev=sv.rev;cho=sv.cho;cur=sv.cur||0;selIdx=sv.selIdx||selIdx;return true;}}catch(e){}return false;}
function buildQuestions(idx){qs=[];idx.forEach(i=>shuf(TESTS[i].preguntas).forEach(q=>qs.push({...q})));qs=shuf(qs);qs.forEach(q=>{if(!q._opts)q._opts=shuf([q.correct,...q.wrong]);});ans=qs.map(()=>null);rev=qs.map(()=>false);cho=qs.map(()=>null);cur=0;}
function score(){return ans.filter(a=>a==='correct').length;}
function done(){return rev.filter(Boolean).length;}
function renderSelect(){
  app.innerHTML=`<div class="card sel-screen"><div class="sel-title">${TITULO}</div><div class="sel-sub">Elige los tests que quieres practicar</div><div class="test-list">${TESTS.map((t,i)=>`<label class="test-item on" id="lbl${i}"><input type="checkbox" id="chk${i}" checked onchange="toggleItem(${i})"><div><div class="t-name">${t.nombre}</div><div class="t-count">${t.preguntas.length} preguntas</div></div></label>`).join('')}</div><button class="btn-start" onclick="startSel()">Empezar &rarr;</button></div>`;
}
function toggleItem(i){document.getElementById('lbl'+i).classList.toggle('on',document.getElementById('chk'+i).checked);}
function startSel(){selIdx=TESTS.map((_,i)=>document.getElementById('chk'+i).checked?i:-1).filter(i=>i>=0);if(!selIdx.length){alert('Selecciona al menos un test.');return;}buildQuestions(selIdx);save();renderQuiz();}
function renderQuiz(){
  const q=qs[cur],opts=q._opts,ci=opts.indexOf(q.correct),isRev=rev[cur];
  const imgHtml=q.img?`<img class="qimg" src="data:image/jpeg;base64,${q.img}" alt="">`:'';
  const optsHtml=opts.map((o,i)=>{let c='opt';if(isRev){if(i===ci)c+=' ok';else if(i===cho[cur])c+=' bad';}else if(i===cho[cur])c+=' sel';return`<button class="${c}"${isRev?' disabled':''} onclick="pick(${i})">${o}</button>`;}).join('');
  const navHtml=isRev
    ?`<button class="btn btn-prev" onclick="prev()"${cur===0?' disabled':''}>&larr;</button><span style="flex:1"></span><button class="btn btn-next" onclick="next()">${cur<qs.length-1?'Siguiente &rarr;':'Ver resultado &rarr;'}</button>`
    :`<button class="btn btn-prev" onclick="prev()"${cur===0?' disabled':''}>&larr;</button><button class="btn btn-skip" onclick="skipQ()">Saltar</button><button class="btn btn-show" onclick="reveal()">Mostrar<span class="kbd">Esp</span></button><button class="btn btn-next" onclick="next()"${cur>=qs.length-1?' disabled':''}>&rarr;</button>`;
  const gridHtml=qs.map((_,i)=>{let c='gq';if(ans[i]==='correct')c+=' ok';else if(ans[i]==='wrong')c+=' bad';if(i===cur)c+=' cur';return`<button class="${c}" onclick="goTo(${i})">${i+1}</button>`;}).join('');
  const pct=qs.length?Math.round(cur/qs.length*100):0;
  app.innerHTML=`<div class="hdr"><span class="hdr-title">${TITULO}</span><span class="hdr-score">${score()}/${done()}</span></div><div class="pbar"><div class="pfill" style="width:${pct}%"></div></div><div class="main"><div class="quiz-area"><div class="card"><div class="qnum">Pregunta ${cur+1} / ${qs.length}</div><div class="qtext">${q.q}</div>${imgHtml}<div class="opts">${optsHtml}</div><div class="navbar">${navHtml}</div></div></div><div class="grid-panel"><div class="gp-title">Preguntas</div><div class="grid">${gridHtml}</div><div class="legend"><div class="li"><div class="ld" style="background:#0f172a;outline:1px solid #334155"></div>Sin responder</div><div class="li"><div class="ld" style="background:#15803d"></div>Correcta</div><div class="li"><div class="ld" style="background:#b91c1c"></div>Incorrecta</div></div>${MULTI?`<button class="btn-change" onclick="renderSelect()">Cambiar tests</button>`:''}</div></div>`;
}
function pick(i){if(rev[cur])return;cho[cur]=i;renderQuiz();}
function reveal(){if(rev[cur])return;rev[cur]=true;const ci=qs[cur]._opts.indexOf(qs[cur].correct);ans[cur]=(cho[cur]===ci)?'correct':'wrong';save();renderQuiz();}
function prev(){if(cur>0){cur--;renderQuiz();}}
function next(){if(cur<qs.length-1){cur++;renderQuiz();}else if(done()===qs.length)renderEnd();}
function goTo(i){cur=i;renderQuiz();}
function skipQ(){for(let i=cur+1;i<qs.length;i++){if(!rev[i]){cur=i;renderQuiz();return;}}for(let i=0;i<cur;i++){if(!rev[i]){cur=i;renderQuiz();return;}}renderEnd();}
function renderEnd(){const tot=done(),sc=score(),p=tot?Math.round(sc/tot*100):0;app.innerHTML=`<div class="card end"><div class="end-pct">${p}%</div><div class="end-sub">${sc} de ${tot} respondidas correctamente</div><div class="end-btns"><button class="btn-restart" onclick="restart()">&#x1F504; De nuevo</button>${MULTI?`<button class="btn-restart" style="background:#334155" onclick="renderSelect()">Cambiar tests</button>`:''}</div></div>`;}
function restart(){buildQuestions(selIdx);save();renderQuiz();}
document.addEventListener('keydown',e=>{
  if(['INPUT','TEXTAREA'].includes(document.activeElement.tagName))return;
  if(e.code==='Space'){e.preventDefault();if(!rev[cur])reveal();}
  if(e.code==='ArrowLeft')prev();
  if(e.code==='ArrowRight'&&rev[cur])next();
  if(e.key==='s'||e.key==='S')skipQ();
  const n=parseInt(e.key);if(!isNaN(n)&&n>=1&&n<=9&&!rev[cur])pick(n-1);
});
(function(){if(MULTI&&!loadSaved()){renderSelect();}else{if(!loadSaved())buildQuestions(selIdx);renderQuiz();}})();
"""

_QUIZ_HTML = (
    "<!DOCTYPE html><html lang='es'><head>"
    "<meta charset='UTF-8'><meta name='viewport' content='width=device-width,initial-scale=1'>"
    "<title>QUIZ_TITLE</title><style>" + _QUIZ_CSS + "</style></head>"
    "<body><div id='app'></div><script>" + _QUIZ_JS + "</script></body></html>"
)


def generar_html_quiz(tests_datos: list[dict], nombre: str) -> bytes:
    tests_js = []
    for test in tests_datos:
        preguntas = []
        for p in test["preguntas"]:
            correcta = None
            incorrectas: list[str] = []
            for texto, es_correcta in p["opciones"]:
                if not texto:
                    continue
                if es_correcta:
                    correcta = texto
                else:
                    incorrectas.append(texto)
            if correcta is None:
                continue
            img_b64 = ""
            if p.get("img_bytes"):
                img_b64 = base64.b64encode(p["img_bytes"]).decode()
            preguntas.append({"q": p["enunciado"], "img": img_b64,
                               "correct": correcta, "wrong": incorrectas})
        if preguntas:
            tests_js.append({"nombre": test["titulo"], "preguntas": preguntas})

    data_json = json.dumps(tests_js, ensure_ascii=False).replace("</script>", "<\\/script>")
    nombre_safe = nombre.replace('"', '').replace("'", "")
    html = _QUIZ_HTML.replace("QUIZ_TESTS_DATA", data_json).replace("QUIZ_TITLE", nombre_safe)
    return html.encode("utf-8")


# ---------------------------------------------------------------------------
# Gemini AI
# ---------------------------------------------------------------------------

_PROMPT_GEMINI = """
Extrae TODAS las preguntas de opcion multiple de TODO el contenido proporcionado \
(puede haber texto, imagenes, PDFs y documentos combinados).
Combina todas las fuentes en un unico conjunto de preguntas.
Devuelve SOLO JSON valido, sin ningun texto adicional, con esta estructura exacta:
{
  "titulo": "titulo descriptivo del examen o test",
  "preguntas": [
    {
      "enunciado": "texto completo y limpio de la pregunta",
      "correcta": "texto exacto de la opcion correcta sin ninguna marca",
      "incorrectas": ["opcion incorrecta 1", "opcion incorrecta 2", "opcion incorrecta 3"]
    }
  ]
}
Reglas:
- Limpia artefactos OCR, mal encoding y caracteres extraños
- La opcion correcta puede estar marcada con asterisco (*), letra circulada, subrayado, (C), flecha, o cualquier marca. Identifícala y elimina la marca del texto
- Si no puedes identificar la correcta con seguridad, omite esa pregunta
- Extrae absolutamente TODAS las preguntas de TODAS las fuentes, sin excepcion
- El titulo debe ser descriptivo (ej: "Traumatologia Tema 1", "Bioquimica Parcial 2024")
"""


def _texto_de_docx(datos: bytes) -> str:
    """Extrae texto plano de un .docx."""
    doc = Document(BytesIO(datos))
    lineas = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            lineas.append("  |  ".join(c.text for c in fila.cells if c.text.strip()))
    return "\n".join(lineas)


_MODELOS_GEMINI = [
    "gemini-2.0-flash",
    "gemini-2.5-flash-preview-05-20",
    "gemini-2.0-flash-lite",
    "gemini-1.5-flash-latest",
    "gemini-1.5-pro-latest",
]

_PROMPT_CORRECCION = """\
Aqui tienes el JSON actual con las preguntas extraidas:
{json_actual}

El usuario pide esta correccion o ajuste:
"{peticion}"

Aplica los cambios y devuelve el JSON COMPLETO actualizado con la misma estructura exacta.
SOLO devuelve JSON valido, sin ningun texto adicional.
"""


def listar_modelos_gemini(api_key: str) -> list[str]:
    """
    Consulta ListModels (NO consume cuota de generación) y devuelve los
    nombres de modelos que soportan generateContent, disponibles para
    esta API key concreta. Lanza excepción si la key es inválida.
    """
    headers = {"x-goog-api-key": api_key}
    modelos: list[str] = []
    ultimo_error = None
    for version in ("v1beta", "v1"):
        url = f"https://generativelanguage.googleapis.com/{version}/models"
        try:
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            for m in resp.json().get("models", []):
                if "generateContent" in m.get("supportedGenerationMethods", []):
                    nombre = m["name"].replace("models/", "")
                    if nombre not in modelos:
                        modelos.append(nombre)
            if modelos:
                break
        except Exception as e:
            ultimo_error = e
            continue
    if not modelos and ultimo_error:
        raise ultimo_error
    # Ordenar: flash primero (más rápidos y con más cuota), luego el resto
    def _orden(n: str):
        return (
            0 if "flash" in n and "lite" not in n else
            1 if "lite" in n else
            2 if "flash" in n else 3,
            n,
        )
    return sorted(modelos, key=_orden)


def diagnosticar_modelo(api_key: str, modelo: str) -> tuple[str, str]:
    """
    Hace UNA petición mínima ('ping') a un modelo para saber si funciona.
    Devuelve (estado, detalle): estado ∈ {"ok", "cuota", "no_existe", "error"}.
    Consume 1 request de cuota SOLO si el modelo está disponible (si da 429
    o 404 no cuenta como uso).
    """
    headers = {"x-goog-api-key": api_key, "Content-Type": "application/json"}
    body = {"contents": [{"parts": [{"text": "ping"}]}],
            "generationConfig": {"maxOutputTokens": 1}}
    ultimo = ("error", "sin respuesta")
    for version in ("v1", "v1beta"):
        url = f"https://generativelanguage.googleapis.com/{version}/models/{modelo}:generateContent"
        try:
            resp = requests.post(url, headers=headers, json=body, timeout=30)
        except Exception as e:
            ultimo = ("error", str(e)[:80])
            continue
        if resp.status_code == 200:
            return ("ok", "Funciona ✓")
        if resp.status_code == 404:
            ultimo = ("no_existe", "no existe en esta API")
            continue
        if resp.status_code == 429:
            qv = None
            try:
                for d in resp.json().get("error", {}).get("details", []):
                    if "QuotaFailure" in d.get("@type", ""):
                        for v in d.get("violations", []):
                            qv = v.get("quotaValue", qv)
            except Exception:
                pass
            if qv is not None and str(qv) == "0":
                return ("cuota", "cuota CERO (no disponible en tu tier)")
            return ("cuota", f"cuota agotada (límite {qv if qv is not None else '?'})")
        try:
            msg = resp.json().get("error", {}).get("message", "")[:80]
        except Exception:
            msg = resp.text[:80]
        ultimo = ("error", f"HTTP {resp.status_code}: {msg}")
    return ultimo


def _gemini_call(api_key: str, modelo: str, partes: list, log=None) -> str:
    """
    Llama al REST API de Gemini con retry en 429.
    - Intenta v1 primero; si devuelve 404, reintenta con v1beta.
    - En caso de 429: reintenta con backoff exponencial (5-60s), máx 5 intentos.
    - log: callable opcional que recibe mensajes de progreso en tiempo real.
    partes: lista de str (texto) o dict {"mime_type": str, "data": bytes}
    """
    import time

    def _log(msg: str):
        if log:
            log(msg)

    headers = {"x-goog-api-key": api_key, "Content-Type": "application/json"}

    parts_json = []
    n_texto = n_img = n_pdf = 0
    for parte in partes:
        if isinstance(parte, str):
            parts_json.append({"text": parte})
            n_texto += 1
        elif isinstance(parte, dict) and "data" in parte:
            mime = parte["mime_type"]
            parts_json.append({
                "inlineData": {
                    "mimeType": mime,
                    "data": base64.b64encode(parte["data"]).decode("utf-8"),
                }
            })
            if mime == "application/pdf":
                n_pdf += 1
            else:
                n_img += 1

    tam_kb = sum(len(p.get("inlineData", {}).get("data", "")) for p in parts_json) // 1024
    _log(f"📦 Preparado: {n_texto} bloque(s) de texto, {n_img} imagen(es), "
         f"{n_pdf} PDF(s) · {tam_kb} KB en base64")

    body = {"contents": [{"parts": parts_json}]}
    esperas = [5, 10, 20, 40, 60]  # backoff entre reintentos por 429

    for version in ("v1", "v1beta"):
        url = (
            f"https://generativelanguage.googleapis.com/{version}"
            f"/models/{modelo}:generateContent"
        )
        for intento in range(len(esperas)):
            _log(f"🌐 [{version}] Enviando petición a Gemini "
                 f"(intento {intento + 1}/{len(esperas)})...")
            t0 = time.time()
            try:
                resp = requests.post(url, headers=headers, json=body, timeout=120)
            except requests.Timeout:
                _log(f"⏱️ Timeout tras 120s en {version}. Reintentando...")
                continue
            dt = time.time() - t0

            if resp.status_code == 404:
                _log(f"↪️ [{version}] Modelo no disponible (404). "
                     f"Probando siguiente versión de API...")
                break  # probar siguiente version de API

            if resp.status_code == 429:
                # Extraer el mensaje y detalles reales de Google
                msg_api = ""
                es_diario = False
                quota_value = None
                quota_id = ""
                try:
                    error_obj = resp.json().get("error", {})
                    msg_api = error_obj.get("message", "")
                    for detalle in error_obj.get("details", []):
                        tipo = detalle.get("@type", "")
                        if "QuotaFailure" in tipo:
                            for v in detalle.get("violations", []):
                                quota_id = v.get("quotaId", "") or quota_id
                                quota_value = v.get("quotaValue", quota_value)
                                qid = quota_id + v.get("quotaMetric", "")
                                if "PerDay" in qid or "per_day" in qid.lower():
                                    es_diario = True
                except Exception:
                    msg_api = resp.text[:300]

                if quota_id:
                    _log(f"📊 Límite alcanzado: `{quota_id}` "
                         f"(valor: {quota_value if quota_value is not None else '?'})")
                if msg_api:
                    _log(f"📋 Google dice: {msg_api[:250]}")

                # quotaValue == 0 → el modelo NO está disponible en este tier
                if quota_value is not None and str(quota_value) == "0":
                    _log("🚫 Este modelo tiene cuota CERO en tu tier — no está disponible.")
                    raise requests.HTTPError(
                        f"429-CERO — El modelo '{modelo}' tiene límite 0 en tu plan "
                        f"(quota '{quota_id}'). NO es que lo hayas agotado: simplemente no está "
                        "disponible en el tier gratuito para tu cuenta/región. "
                        "Prueba OTRO modelo (pulsa Verificar key y prueba el Diagnóstico) "
                        "o activa facturación.",
                        response=resp,
                    )

                # Si es límite DIARIO, reintentar no sirve de nada
                if es_diario:
                    _log("🚫 Límite DIARIO agotado — reintentar no ayudará hoy.")
                    raise requests.HTTPError(
                        f"429-DIARIO — Cuota DIARIA agotada (quota '{quota_id}', "
                        f"límite {quota_value}). {msg_api[:150]} "
                        "Soluciones: prueba otro modelo, espera 24h, o activa facturación.",
                        response=resp,
                    )

                if intento < len(esperas) - 1:
                    espera = esperas[intento]
                    _log(f"⏳ Rate limit por minuto (429). Esperando {espera}s antes de reintentar...")
                    time.sleep(espera)
                    continue
                raise requests.HTTPError(
                    f"429 — Rate limit tras varios intentos. {msg_api[:200]} "
                    "Espera 1-2 minutos y reintenta.",
                    response=resp,
                )

            if resp.status_code != 200:
                # Mostrar el mensaje de error real de la API
                try:
                    err = resp.json().get("error", {}).get("message", resp.text[:200])
                except Exception:
                    err = resp.text[:200]
                _log(f"❌ [{version}] HTTP {resp.status_code}: {err}")
                resp.raise_for_status()

            _log(f"✅ Respuesta recibida en {dt:.1f}s. Procesando JSON...")
            return resp.json()["candidates"][0]["content"]["parts"][0]["text"]

    raise requests.HTTPError(
        f"404 — El modelo '{modelo}' no existe en las APIs v1 ni v1beta. "
        "Prueba con 'gemini-2.0-flash' o 'gemini-2.5-flash-preview-05-20'."
    )


def aplicar_correccion_gemini(api_key: str, modelo: str,
                               datos_actuales: dict, peticion: str, log=None) -> dict:
    prompt = _PROMPT_CORRECCION.format(
        json_actual=json.dumps(datos_actuales, ensure_ascii=False, indent=2),
        peticion=peticion,
    )
    raw = _gemini_call(api_key, modelo, [prompt], log=log).strip()
    for marker in ["```json", "```"]:
        if marker in raw:
            raw = raw.split(marker, 1)[1].split("```")[0].strip()
            break
    return json.loads(raw)


def extraer_con_gemini_multi(api_key: str, modelo: str = "gemini-2.0-flash",
                              texto: str = "", archivos: list | None = None,
                              log=None) -> dict:
    """
    Llama directamente al REST API de Gemini v1 (sin SDK de Google).
    Los .docx se convierten a texto; PDFs e imágenes se envían como
    inline_data en base64 — formato nativo de la API REST.
    """
    def _log(msg: str):
        if log:
            log(msg)

    partes: list = [_PROMPT_GEMINI]
    texto_acumulado = texto.strip()

    for archivo in (archivos or []):
        datos = archivo.read()
        ext = archivo.name.rsplit(".", 1)[-1].lower()
        if ext in ("docx", "doc"):
            _log(f"📄 Extrayendo texto de {archivo.name}...")
            try:
                texto_doc = _texto_de_docx(datos)
                texto_acumulado += f"\n\n--- {archivo.name} ---\n{texto_doc}"
                _log(f"   → {len(texto_doc)} caracteres extraídos de {archivo.name}")
            except Exception as e:
                _log(f"   ⚠️ No se pudo leer {archivo.name}: {e}")
        elif ext == "pdf":
            _log(f"📎 Adjuntando PDF {archivo.name} ({len(datos)//1024} KB)...")
            partes.append({"mime_type": "application/pdf", "data": datos})
        elif ext in ("jpg", "jpeg", "png", "webp"):
            _log(f"🖼️ Adjuntando imagen {archivo.name} ({len(datos)//1024} KB)...")
            mime = archivo.type or f"image/{ext.replace('jpg', 'jpeg')}"
            partes.append({"mime_type": mime, "data": datos})

    if texto_acumulado:
        partes.append(f"\n\n--- TEXTO / DOCUMENTOS ---\n{texto_acumulado}")

    if len(partes) == 1:
        raise ValueError("No hay contenido para procesar.")

    raw = _gemini_call(api_key, modelo, partes, log=log).strip()
    _log("🔍 Parseando preguntas del JSON devuelto...")
    for marker in ["```json", "```"]:
        if marker in raw:
            raw = raw.split(marker, 1)[1].split("```")[0].strip()
            break

    return json.loads(raw)


def gemini_a_tests(datos: dict) -> list[dict]:
    """Convierte la salida JSON de Gemini al formato interno de todos_los_tests."""
    titulo = datos.get("titulo", "Test importado con IA")
    preguntas_con_img: list[dict] = []
    for p in datos.get("preguntas", []):
        correcta = (p.get("correcta") or "").strip()
        if not correcta:
            continue
        incorrectas = [i.strip() for i in p.get("incorrectas", []) if i.strip()]
        opciones = [(correcta, True)] + [(inc, False) for inc in incorrectas]
        preguntas_con_img.append({
            "enunciado": (p.get("enunciado") or "").strip(),
            "opciones": opciones,
            "img_bytes": None,
            "img_nombre": None,
        })

    doc = Document()
    doc.add_heading(titulo, 1)
    for i, p in enumerate(preguntas_con_img, 1):
        procesar_pregunta(doc, i, p["enunciado"], None, p["opciones"])
    buf = BytesIO()
    doc.save(buf)

    return [{"titulo": titulo, "preguntas": preguntas_con_img, "_word_bytes": buf.getvalue()}]


def construir_resultado(todos_los_tests: list[dict], errores: list[str] | None = None) -> dict:
    """Genera todos los exportables y los empaqueta en el dict de session_state."""
    nb = generar_nombre_base(todos_los_tests)
    es_multi = len(todos_los_tests) > 1

    return {
        "nombre_base": nb,
        "es_multi": es_multi,
        "word_combinado_con": generar_word_combinado(todos_los_tests, con_respuesta=True),
        "word_combinado_sin": generar_word_combinado(todos_los_tests, con_respuesta=False),
        "zip_word_individuales": generar_zip_word_individuales(todos_los_tests) if es_multi else b"",
        "pdf_combinado_con": generar_pdf(todos_los_tests, con_respuesta=True),
        "pdf_combinado_sin": generar_pdf(todos_los_tests, con_respuesta=False),
        "zip_pdf_individuales": generar_zip_pdf_individuales(todos_los_tests) if es_multi else b"",
        "zip_imagenes": generar_zip_imagenes(todos_los_tests),
        "zip_remnote": generar_zip_remnote(todos_los_tests),
        "apkg_anki": generar_apkg_anki(todos_los_tests),
        "html_quiz": generar_html_quiz(todos_los_tests, nb),
        "tests_ok": len(todos_los_tests),
        "errores": errores or [],
    }


# ---------------------------------------------------------------------------
# Interfaz de usuario (Streamlit)
# ---------------------------------------------------------------------------

st.title("📝 Daypo Extractor")
st.markdown(
    "Extrae y exporta preguntas de [Daypo](https://www.daypo.com) — o importa "
    "cualquier examen con IA — a Word, PDF, RemNote, Anki y quiz offline."
)
st.divider()

for _k, _v in [("resultado", None), ("ia_estado", None),
               ("ia_datos", None), ("ia_chat", []),
               ("ia_api_key", ""), ("ia_modelo", "gemini-2.0-flash"),
               ("ia_modelos_disponibles", None),
               ("ia_requests_sesion", 0), ("ia_diagnostico", None)]:
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ── Pantalla de resultados (compartida por ambas pestañas) ─────────────────
if st.session_state["resultado"] is not None:
    res = st.session_state["resultado"]

    if res["errores"]:
        st.warning(f"No se pudieron procesar {len(res['errores'])} enlace(s).")

    st.success(f"Procesados {res['tests_ok']} test(s). Archivos listos para descargar.")

    nb = res["nombre_base"]
    es_multi = res["es_multi"]

    # Selector de version para Word y PDF
    version = st.radio("Versión Word/PDF:", ["Con respuestas", "Sin respuestas"],
                       horizontal=True)
    suf = "con" if version == "Con respuestas" else "sin"

    # Fila 1
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            label="📄 Word" + (" (todos)" if es_multi else ""),
            data=res[f"word_combinado_{suf}"],
            file_name=f"{nb}_Word_{suf}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with c2:
        st.download_button(
            label="🧠 RemNote MCQ",
            data=res["zip_remnote"],
            file_name=f"{nb}_RemNote.zip",
            mime="application/zip",
            use_container_width=True,
            help="Importa el ZIP en RemNote: ajustes → Importar → Markdown.",
        )
    with c3:
        st.download_button(
            label="🃏 Anki (.apkg)",
            data=res["apkg_anki"],
            file_name=f"{nb}_Anki.apkg",
            mime="application/octet-stream",
            use_container_width=True,
            help="Doble clic en el .apkg para importar en Anki.",
        )

    # Fila 2
    c4, c5, c6 = st.columns(3)
    with c4:
        st.download_button(
            label="🖨️ PDF" + (" (todos)" if es_multi else ""),
            data=res[f"pdf_combinado_{suf}"],
            file_name=f"{nb}_{suf}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    with c5:
        st.download_button(
            label="🌐 Mini-App Quiz",
            data=res["html_quiz"],
            file_name=f"{nb}_Quiz.html",
            mime="text/html",
            use_container_width=True,
            help="Abre en cualquier navegador. Funciona offline.",
        )
    with c6:
        st.download_button(
            label="🖼️ Imagenes sueltas",
            data=res["zip_imagenes"],
            file_name=f"{nb}_Imagenes.zip",
            mime="application/zip",
            use_container_width=True,
        )

    # Botones "por separado" (solo con varios tests)
    if es_multi:
        st.markdown("---")
        ca, cb = st.columns(2)
        with ca:
            st.download_button(
                label="📄 Word por separado (ZIP)",
                data=res["zip_word_individuales"],
                file_name=f"{nb}_Word_individual.zip",
                mime="application/zip",
                use_container_width=True,
                help="ZIP con Con_Respuesta/ y Sin_Respuesta/ — un .docx por test.",
            )
        with cb:
            st.download_button(
                label="🖨️ PDF por separado (ZIP)",
                data=res["zip_pdf_individuales"],
                file_name=f"{nb}_PDF_individual.zip",
                mime="application/zip",
                use_container_width=True,
                help="ZIP con Con_Respuesta/ y Sin_Respuesta/ — un .pdf por test.",
            )

    with st.expander("ℹ️ Instrucciones de importacion"):
        st.markdown("""
**RemNote:** descarga el ZIP → RemNote → ajustes → Importar → Markdown → sube el ZIP.

**Anki:** doble clic en el .apkg (o Archivo → Importar). Las opciones aparecen en orden aleatorio en cada repaso; clic para seleccionar, Espacio para revelar, el orden no cambia al girar la carta.
""")

    st.divider()
    if st.button("🔄 Nueva extraccion / importacion", use_container_width=True):
        st.session_state["resultado"] = None
        st.rerun()

    st.stop()

# ── Pestañas de entrada ────────────────────────────────────────────────────
tab_daypo, tab_ia = st.tabs(["🔗 Extraer de Daypo", "🤖 Importar con IA (Gemini)"])

# ── Pestaña 1: Daypo ──────────────────────────────────────────────────────
with tab_daypo:
    enlaces_texto = st.text_area(
        "Pega aqui tus enlaces de Daypo (o cualquier texto que los contenga):",
        height=200,
        placeholder="Pega aqui tus enlaces de Daypo...",
    )

    if enlaces_texto.strip():
        detectados = extraer_enlaces_daypo(enlaces_texto)
        if detectados:
            st.success(f"🔍 {len(detectados)} enlace(s) detectado(s)")
            for e in detectados:
                st.caption(e)
        else:
            st.warning("No se han detectado enlaces de Daypo en el texto introducido.")

    col1, col2 = st.columns([3, 1])
    with col1:
        iniciar = st.button("Iniciar Extraccion", type="primary", use_container_width=True)
    with col2:
        st.caption("Requiere internet")

    if iniciar:
        enlaces = extraer_enlaces_daypo(enlaces_texto)
        if not enlaces:
            st.error("No se han detectado enlaces validos de Daypo.")
            st.stop()

        headers_web = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        todos_los_tests: list[dict] = []
        barra = st.progress(0, text="Iniciando...")
        log = st.empty()
        errores = []

        for indice, url in enumerate(enlaces):
            nombre_corto = url.rstrip("/").split("/")[-1].replace(".html", "")
            barra.progress(int((indice / len(enlaces)) * 100),
                           text=f"Procesando {indice + 1}/{len(enlaces)}: {nombre_corto}")
            log.info(f"Procesando: {url}")

            resultado = extraer_datos_test(url, headers_web)
            if resultado is None:
                errores.append(url)
                st.warning(f"No se pudo procesar: {url}")
                continue

            id_test, titulo, preguntas = resultado
            doc_indiv = Document()
            doc_indiv.add_heading(titulo, 1)
            preguntas_con_img: list[dict] = []

            for i, pregunta in enumerate(preguntas, start=1):
                img_bytes = None
                img_nombre = None
                if pregunta["num_imagen"] is not None:
                    img_nombre = f"img_{id_test}_{pregunta['num_imagen']}.jpg"
                    img_bytes = obtener_imagen(id_test, pregunta["num_imagen"], url)
                procesar_pregunta(doc_indiv, i, pregunta["enunciado"], img_bytes,
                                  pregunta["opciones"], con_respuesta=True)
                preguntas_con_img.append({
                    "enunciado": pregunta["enunciado"],
                    "opciones": pregunta["opciones"],
                    "img_bytes": img_bytes,
                    "img_nombre": img_nombre,
                })

            buf_indiv = BytesIO()
            doc_indiv.save(buf_indiv)
            todos_los_tests.append({
                "titulo": titulo,
                "preguntas": preguntas_con_img,
                "_word_bytes": buf_indiv.getvalue(),
            })

        barra.progress(100, text="Extraccion completada.")
        log.empty()

        resultado_final = construir_resultado(todos_los_tests, errores=errores)
        resultado_final["tests_ok"] = len(todos_los_tests)
        st.session_state["resultado"] = resultado_final
        st.rerun()

# ── Pestaña 2: Gemini AI ──────────────────────────────────────────────────
with tab_ia:

    # ── Vista de correcciones (tras extracción) ───────────────────────────
    if st.session_state["ia_estado"] == "preview":
        datos = st.session_state["ia_datos"]
        api_key_chat = st.session_state["ia_api_key"]
        modelo_chat = st.session_state["ia_modelo"]
        n_q = len(datos.get("preguntas", []))

        st.success(f"**{n_q} preguntas extraídas** — *{datos.get('titulo', '')}*")

        with st.expander("Ver preguntas extraídas", expanded=False):
            for i, p in enumerate(datos.get("preguntas", []), 1):
                st.markdown(f"**{i}. {p['enunciado']}**")
                for inc in p.get("incorrectas", []):
                    st.caption(f"  × {inc}")
                st.caption(f"  ✓ **{p.get('correcta', '')}**")
                if i < n_q:
                    st.divider()

        st.markdown("---")
        st.markdown("**Correcciones con IA** — describe cualquier ajuste y Gemini actualiza las preguntas:")

        # Historial de chat
        for msg in st.session_state["ia_chat"]:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        # Input de corrección
        correccion = st.chat_input(
            "Ej: 'La correcta de la pregunta 3 es la opción B' · 'Elimina la pregunta 7' · 'Cambia el título a...'"
        )
        if correccion:
            st.session_state["ia_chat"].append({"role": "user", "content": correccion})
            respuesta = None
            with st.status("Aplicando corrección...", expanded=True) as estado:
                def _log_chat(msg: str):
                    estado.write(msg)
                try:
                    nuevos_datos = aplicar_correccion_gemini(
                        api_key_chat, modelo_chat, datos, correccion, log=_log_chat
                    )
                    st.session_state["ia_requests_sesion"] += 1
                    st.session_state["ia_datos"] = nuevos_datos
                    n_nuevo = len(nuevos_datos.get("preguntas", []))
                    respuesta = f"Listo. Ahora hay **{n_nuevo} preguntas** en *{nuevos_datos.get('titulo', '')}*."
                    estado.update(label="Corrección aplicada", state="complete")
                except json.JSONDecodeError:
                    respuesta = "No pude aplicar la corrección (Gemini no devolvió JSON válido). Inténtalo de otra forma."
                    estado.update(label="JSON inválido", state="error")
                except requests.HTTPError as e:
                    estado.update(label="Error de Gemini", state="error")
                    if "429" in str(e):
                        respuesta = (
                            "⏳ **Rate limit** — el tier gratuito sigue al límite (~15 req/min). "
                            "Espera 1-2 minutos y vuelve a intentarlo."
                        )
                    else:
                        respuesta = f"Error de Gemini: {str(e)[:100]}"
                except Exception as e:
                    respuesta = f"Error: {str(e)[:200]}"
                    estado.update(label="Error", state="error")

            if respuesta:
                st.session_state["ia_chat"].append({"role": "assistant", "content": respuesta})
                st.rerun()

        st.markdown("---")
        col_gen, col_volver = st.columns(2)
        with col_gen:
            if st.button("✅ Generar todos los archivos", type="primary", use_container_width=True):
                todos_ia = gemini_a_tests(st.session_state["ia_datos"])
                st.session_state["resultado"] = construir_resultado(todos_ia)
                st.session_state["ia_estado"] = None
                st.session_state["ia_chat"] = []
                st.rerun()
        with col_volver:
            if st.button("↩️ Empezar de nuevo", use_container_width=True):
                st.session_state["ia_estado"] = None
                st.session_state["ia_datos"] = None
                st.session_state["ia_chat"] = []
                st.rerun()

    # ── Formulario de entrada ─────────────────────────────────────────────
    else:
        st.subheader("Importar preguntas con IA")
        st.markdown(
            "Combina texto pegado, fotos de exámenes, PDFs y documentos Word "
            "en una sola llamada — Gemini extrae y unifica todas las preguntas."
        )

        api_key = st.text_input(
            "Clave API de Gemini:",
            type="password",
            value=st.session_state["ia_api_key"],
            help="Clave gratuita en [aistudio.google.com](https://aistudio.google.com)",
        )

        col_verif, col_diag = st.columns(2)
        with col_verif:
            verificar = st.button("🔍 Verificar key", use_container_width=True,
                                  disabled=not bool(api_key))
        with col_diag:
            diagnostico = st.button("🩺 Diagnóstico de modelos", use_container_width=True,
                                    disabled=not bool(api_key))

        if verificar:
            with st.spinner("Consultando modelos disponibles para tu API key..."):
                try:
                    modelos = listar_modelos_gemini(api_key)
                    if modelos:
                        st.session_state["ia_modelos_disponibles"] = modelos
                        st.session_state["ia_api_key"] = api_key
                        st.success(f"✅ API key válida. {len(modelos)} modelos disponibles.")
                    else:
                        st.warning("La key es válida pero no devolvió modelos con generateContent.")
                except requests.HTTPError as e:
                    code = getattr(e.response, "status_code", "?")
                    st.error(f"❌ API key rechazada (HTTP {code}). Revisa que la copiaste completa "
                             "y que está activa en aistudio.google.com.")
                except Exception as e:
                    st.error(f"❌ No se pudo verificar la key: {str(e)[:200]}")

        if diagnostico:
            st.session_state["ia_api_key"] = api_key
            modelos_a_probar = st.session_state["ia_modelos_disponibles"] or _MODELOS_GEMINI
            resultados = []
            with st.status("Probando cada modelo con un ping mínimo...", expanded=True) as est:
                for m in modelos_a_probar:
                    est.write(f"⏳ Probando `{m}`...")
                    estado_m, detalle_m = diagnosticar_modelo(api_key, m)
                    if estado_m == "ok":
                        st.session_state["ia_requests_sesion"] += 1
                    icono = {"ok": "✅", "cuota": "🚫", "no_existe": "❔", "error": "⚠️"}.get(estado_m, "•")
                    est.write(f"{icono} `{m}` — {detalle_m}")
                    resultados.append((m, estado_m, detalle_m))
                est.update(label="Diagnóstico completado", state="complete")
            st.session_state["ia_diagnostico"] = resultados

        # Mostrar resultados del último diagnóstico
        if st.session_state["ia_diagnostico"]:
            ok_models = [m for m, e, _ in st.session_state["ia_diagnostico"] if e == "ok"]
            if ok_models:
                st.success(f"✅ Modelos que funcionan AHORA: {', '.join(f'`{m}`' for m in ok_models)}")
            else:
                st.error(
                    "🚫 **Ningún modelo respondió OK.** Si todos dan 'cuota CERO', tu cuenta no "
                    "tiene acceso al tier gratuito de la API (común en algunas regiones/cuentas "
                    "nuevas). Necesitas **activar facturación** en aistudio.google.com — el pago "
                    "por uso es céntimos por examen."
                )

        # Lista de modelos: dinámica si se verificó, si no la fija como fallback
        opciones_modelo = st.session_state["ia_modelos_disponibles"] or _MODELOS_GEMINI
        modelo_sel = st.selectbox(
            "Modelo:",
            options=opciones_modelo,
            index=(opciones_modelo.index(st.session_state["ia_modelo"])
                   if st.session_state["ia_modelo"] in opciones_modelo else 0),
        )

        # Contador de uso de la sesión
        col_info, col_reset = st.columns([3, 1])
        with col_info:
            if st.session_state["ia_modelos_disponibles"] is None:
                st.caption("💡 Pulsa **Verificar key** para cargar tus modelos, o **Diagnóstico** "
                           "para ver cuáles funcionan ahora mismo.")
            else:
                st.caption(f"🔢 Peticiones de generación hechas en esta sesión: "
                           f"**{st.session_state['ia_requests_sesion']}** "
                           f"· La cuota se ve en [ai.dev/rate-limit](https://ai.dev/rate-limit)")

        texto_input = st.text_area(
            "Texto (opcional) — pega preguntas directamente aquí:",
            height=160,
            placeholder="1. ¿Cuál es el tratamiento?\n* Opción A  ← correcta\nOpción B\nOpción C",
        )

        archivos = st.file_uploader(
            "Archivos (opcional) — imágenes, PDFs o Word (varios a la vez):",
            type=["jpg", "jpeg", "png", "webp", "pdf", "docx"],
            accept_multiple_files=True,
        )

        hay_contenido = bool(texto_input.strip()) or bool(archivos)
        if hay_contenido:
            partes_info = []
            if texto_input.strip():
                partes_info.append(f"📝 Texto ({len(texto_input.strip())} caracteres)")
            for f in (archivos or []):
                ext = f.name.rsplit(".", 1)[-1].upper()
                icono = "🖼️" if ext in ("JPG", "JPEG", "PNG", "WEBP") else "📄"
                partes_info.append(f"{icono} {f.name}")
            st.info("**Contenido a procesar:**\n" + "\n".join(f"- {p}" for p in partes_info))

        procesar_ia = st.button(
            "Procesar con Gemini",
            type="primary",
            use_container_width=True,
            disabled=not (_GEMINI_OK and bool(api_key) and hay_contenido),
        )

        if procesar_ia:
            with st.status("Procesando con Gemini...", expanded=True) as estado:
                mensajes: list[str] = []

                def _log_ui(msg: str):
                    mensajes.append(msg)
                    estado.write(msg)

                try:
                    datos = extraer_con_gemini_multi(
                        api_key, modelo=modelo_sel,
                        texto=texto_input, archivos=archivos,
                        log=_log_ui,
                    )
                    st.session_state["ia_requests_sesion"] += 1
                    if not datos.get("preguntas"):
                        estado.update(label="Sin preguntas encontradas", state="error")
                        st.error("Gemini no encontró preguntas. Prueba con contenido más estructurado.")
                    else:
                        n = len(datos["preguntas"])
                        _log_ui(f"✅ {n} preguntas extraídas.")
                        estado.update(label=f"{n} preguntas extraídas", state="complete")
                        st.session_state["ia_estado"] = "preview"
                        st.session_state["ia_datos"] = datos
                        st.session_state["ia_api_key"] = api_key
                        st.session_state["ia_modelo"] = modelo_sel
                        st.session_state["ia_chat"] = []
                        st.rerun()
                except json.JSONDecodeError:
                    estado.update(label="JSON inválido", state="error")
                    st.error("Gemini no devolvió JSON válido. Prueba dividiendo el contenido.")
                except requests.HTTPError as e:
                    estado.update(label="Error de Gemini", state="error")
                    err_msg = str(e)
                    if "429-CERO" in err_msg:
                        st.error(
                            "🚫 **Este modelo tiene cuota CERO en tu cuenta** — no es que lo hayas "
                            "agotado, simplemente **no está disponible** en tu tier gratuito "
                            "(habitual en cuentas nuevas o ciertas regiones).\n\n"
                            "**Qué hacer:**\n"
                            "1. Pulsa **🩺 Diagnóstico de modelos** para ver cuáles SÍ funcionan.\n"
                            "2. Si ninguno funciona, **activa facturación** en "
                            "[aistudio.google.com](https://aistudio.google.com) (céntimos por examen).\n\n"
                            f"_{err_msg[:250]}_"
                        )
                    elif "429-DIARIO" in err_msg:
                        st.error(
                            "🚫 **Cuota DIARIA agotada** en este modelo del tier gratuito.\n\n"
                            "Reintentar hoy no servirá. Opciones:\n"
                            "1. **Cambia de modelo** (cada uno tiene su propia cuota diaria) — "
                            "usa **🩺 Diagnóstico** para ver cuáles quedan.\n"
                            "2. **Espera a mañana** (la cuota se reinicia cada 24h).\n"
                            "3. **Activa facturación** en [Google AI Studio](https://aistudio.google.com).\n\n"
                            f"_{err_msg.split('429-DIARIO — ')[-1][:300]}_"
                        )
                    elif "429" in err_msg:
                        st.warning(
                            "⏳ **Rate limit por minuto** — Espera 1-2 minutos y reintenta.\n\n"
                            f"_Mensaje de Google: {err_msg[:300]}_"
                        )
                    elif "404" in err_msg:
                        st.error("❌ El modelo seleccionado no existe. Prueba con `gemini-2.0-flash`.")
                    else:
                        st.error(f"Error en Gemini: {err_msg[:400]}")
                except Exception as e:
                    estado.update(label="Error", state="error")
                    st.error(f"Error: {str(e)[:300]}")

        with st.expander("💡 Consejos y solución de problemas"):
            st.markdown("""
**Uso normal:**
- **Todo a la vez:** combina texto + fotos + PDF + Word en una sola operación.
- **PDFs escaneados:** se envían directamente a Gemini para OCR.
- **Word (.docx):** el texto se extrae antes de enviarlo.
- **Correcta:** Gemini detecta asteriscos (*), letras circuladas, (C), subrayados, marcas de lápiz, etc.

**Si recibes error de cuota (429):**
- El tier **gratuito** tiene un límite **diario** de peticiones por modelo (a veces tan bajo como
  ~50 al día). Si lo agotas, **reintentar no sirve hasta el día siguiente**.
- **Cambia de modelo** en el selector: cada modelo tiene su propia cuota diaria independiente.
- Pulsa **🔍 Verificar key** para ver exactamente qué modelos tiene disponibles tu cuenta.
- **Solución definitiva:** activa facturación en
  [aistudio.google.com](https://aistudio.google.com) → es pago por uso y procesar exámenes
  cuesta céntimos, pero los límites suben muchísimo.
- Consulta tu uso actual en [ai.dev/rate-limit](https://ai.dev/rate-limit).
""")
