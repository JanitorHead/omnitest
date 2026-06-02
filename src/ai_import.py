"""Importación de preguntas con IA (Gemini) vía REST API v1/v1beta.

No usa ningún SDK de Google (todos enrutan a v1beta y fallan con modelos
actuales). Llama directamente al endpoint REST con 'requests'.
"""
import base64
import json
import random
import time
from io import BytesIO

import requests
from docx import Document

# Modelos ESTABLES primero (mejor cuota en tier gratuito). Los 'preview'/'exp'
# suelen tener cuota 0 en cuentas gratuitas — se ofrecen pero con advertencia.
# 'flash-lite' va primero: rápido, sin 'thinking' pesado y con cuota más amplia.
MODELOS_GEMINI = [
    "gemini-2.5-flash-lite",
    "gemini-2.5-flash",
    "gemini-flash-latest",
    "gemini-flash-lite-latest",
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
    "gemini-1.5-flash-latest",
    "gemini-2.5-pro",
]

MODELO_POR_DEFECTO = "gemini-2.5-flash-lite"

_PROMPT_GEMINI = """
Extrae TODAS las preguntas de opcion multiple de TODO el contenido proporcionado \
(puede haber texto, imagenes, PDFs y documentos combinados).
Combina todas las fuentes en un unico conjunto de preguntas.
Devuelve SOLO JSON valido, sin ningun texto adicional, con esta estructura exacta:
{
  "titulo": "titulo descriptivo del examen o test",
  "preguntas": [
    {
      "enunciado": "texto completo y limpio de la pregunta",
      "correcta": "texto exacto de la opcion correcta sin ninguna marca",
      "incorrectas": ["opcion incorrecta 1", "opcion incorrecta 2", "opcion incorrecta 3"]
    }
  ]
}
Reglas:
- Limpia artefactos OCR, mal encoding y caracteres extraños
- Extrae absolutamente TODAS las preguntas de TODAS las fuentes, sin excepcion
- IMPORTANTE sobre la respuesta correcta:
  * Si la correcta ESTÁ marcada (asterisco *, letra circulada, subrayado, (C), flecha,
    negrita, color, "respuesta: X", etc.), úsala y elimina la marca del texto.
  * Si NINGUNA opción está marcada como correcta, DEDUCE TÚ la respuesta correcta
    usando tu conocimiento experto sobre la materia. SIEMPRE debes elegir una opción
    como correcta — son preguntas tipo test con una única respuesta válida.
  * NUNCA dejes una pregunta fuera por no estar marcada: razónala y respóndela tú.
- El campo "correcta" SIEMPRE debe contener una de las opciones (nunca vacío).
- El titulo debe ser descriptivo (ej: "Traumatologia Tema 1", "Bioquimica Parcial 2024")
"""

_PROMPT_CORRECCION = """\
Aqui tienes el JSON actual con las preguntas extraidas:
{json_actual}

El usuario pide esta correccion o ajuste:
"{peticion}"

Aplica los cambios y devuelve el JSON COMPLETO actualizado con la misma estructura exacta.
SOLO devuelve JSON valido, sin ningun texto adicional.
"""


def es_preview(modelo: str) -> bool:
    m = modelo.lower()
    return "preview" in m or "-exp" in m or "experimental" in m


def _texto_de_docx(datos: bytes) -> str:
    doc = Document(BytesIO(datos))
    lineas = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            lineas.append("  |  ".join(c.text for c in fila.cells if c.text.strip()))
    return "\n".join(lineas)


def listar_modelos_gemini(api_key: str) -> list[str]:
    """ListModels (no consume cuota). Devuelve modelos con generateContent."""
    headers = {"x-goog-api-key": api_key}
    modelos: list[str] = []
    ultimo_error = None
    for version in ("v1beta", "v1"):
        url = f"https://generativelanguage.googleapis.com/{version}/models"
        try:
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            for m in resp.json().get("models", []):
                if "generateContent" in m.get("supportedGenerationMethods", []):
                    nombre = m["name"].replace("models/", "")
                    if nombre not in modelos:
                        modelos.append(nombre)
            if modelos:
                break
        except Exception as e:  # noqa: BLE001
            ultimo_error = e
            continue
    if not modelos and ultimo_error:
        raise ultimo_error

    def _orden(n: str):
        return (
            0 if "flash" in n and "lite" not in n else
            1 if "lite" in n else
            2 if "flash" in n else 3,
            n,
        )

    return sorted(modelos, key=_orden)


def diagnosticar_modelo(api_key: str, modelo: str) -> tuple[str, str]:
    """Ping mínimo a un modelo. Devuelve (estado, detalle).

    estado ∈ {"ok", "cuota", "no_existe", "error"}. Solo consume 1 request si
    el modelo responde 200.
    """
    headers = {"x-goog-api-key": api_key, "Content-Type": "application/json"}
    body = {"contents": [{"parts": [{"text": "ping"}]}],
            "generationConfig": {"maxOutputTokens": 1}}
    ultimo = ("error", "sin respuesta")
    for version in ("v1", "v1beta"):
        url = f"https://generativelanguage.googleapis.com/{version}/models/{modelo}:generateContent"
        try:
            resp = requests.post(url, headers=headers, json=body, timeout=30)
        except Exception as e:  # noqa: BLE001
            ultimo = ("error", str(e)[:80])
            continue
        if resp.status_code == 200:
            return ("ok", "Funciona ✓")
        if resp.status_code == 404:
            ultimo = ("no_existe", "no existe en esta API")
            continue
        if resp.status_code == 429:
            qv = None
            try:
                for d in resp.json().get("error", {}).get("details", []):
                    if "QuotaFailure" in d.get("@type", ""):
                        for v in d.get("violations", []):
                            qv = v.get("quotaValue", qv)
            except Exception:  # noqa: BLE001
                pass
            if qv is not None and str(qv) == "0":
                return ("cuota", "cuota CERO (no disponible en tu tier)")
            return ("cuota", f"cuota agotada (límite {qv if qv is not None else '?'})")
        try:
            msg = resp.json().get("error", {}).get("message", "")[:80]
        except Exception:  # noqa: BLE001
            msg = resp.text[:80]
        ultimo = ("error", f"HTTP {resp.status_code}: {msg}")
    return ultimo


def _gemini_call(api_key: str, modelo: str, partes: list, log=None) -> str:
    """Llama al REST API de Gemini con fallback v1->v1beta y retry en 429.

    partes: lista de str (texto) o dict {"mime_type": str, "data": bytes}.
    log: callable opcional que recibe mensajes de progreso en tiempo real.
    """
    def _log(msg: str):
        if log:
            log(msg)

    headers = {"x-goog-api-key": api_key, "Content-Type": "application/json"}

    parts_json = []
    n_texto = n_img = n_pdf = 0
    for parte in partes:
        if isinstance(parte, str):
            parts_json.append({"text": parte})
            n_texto += 1
        elif isinstance(parte, dict) and "data" in parte:
            mime = parte["mime_type"]
            parts_json.append({
                "inlineData": {
                    "mimeType": mime,
                    "data": base64.b64encode(parte["data"]).decode("utf-8"),
                }
            })
            if mime == "application/pdf":
                n_pdf += 1
            else:
                n_img += 1

    tam_kb = sum(len(p.get("inlineData", {}).get("data", "")) for p in parts_json) // 1024
    _log(f"📦 Preparado: {n_texto} bloque(s) de texto, {n_img} imagen(es), "
         f"{n_pdf} PDF(s) · {tam_kb} KB en base64")

    body = {"contents": [{"parts": parts_json}]}
    esperas = [5, 10, 20, 40, 60]  # backoff entre reintentos por 429

    for version in ("v1", "v1beta"):
        url = (
            f"https://generativelanguage.googleapis.com/{version}"
            f"/models/{modelo}:generateContent"
        )
        for intento in range(len(esperas)):
            _log(f"🌐 [{version}] Enviando petición a Gemini "
                 f"(intento {intento + 1}/{len(esperas)})...")
            t0 = time.time()
            try:
                resp = requests.post(url, headers=headers, json=body, timeout=120)
            except requests.Timeout:
                _log(f"⏱️ Timeout tras 120s en {version}. Reintentando...")
                continue
            dt = time.time() - t0

            if resp.status_code == 404:
                _log(f"↪️ [{version}] Modelo no disponible (404). "
                     f"Probando siguiente versión de API...")
                break

            if resp.status_code == 429:
                msg_api = ""
                es_diario = False
                quota_value = None
                quota_id = ""
                try:
                    error_obj = resp.json().get("error", {})
                    msg_api = error_obj.get("message", "")
                    for detalle in error_obj.get("details", []):
                        if "QuotaFailure" in detalle.get("@type", ""):
                            for v in detalle.get("violations", []):
                                quota_id = v.get("quotaId", "") or quota_id
                                quota_value = v.get("quotaValue", quota_value)
                                qid = quota_id + v.get("quotaMetric", "")
                                if "PerDay" in qid or "per_day" in qid.lower():
                                    es_diario = True
                except Exception:  # noqa: BLE001
                    msg_api = resp.text[:300]

                if quota_id:
                    _log(f"📊 Límite alcanzado: `{quota_id}` "
                         f"(valor: {quota_value if quota_value is not None else '?'})")
                if msg_api:
                    _log(f"📋 Google dice: {msg_api[:250]}")

                if quota_value is not None and str(quota_value) == "0":
                    _log("🚫 Este modelo tiene cuota CERO en tu tier — no está disponible.")
                    raise requests.HTTPError(
                        f"429-CERO — El modelo '{modelo}' tiene límite 0 en tu plan "
                        f"(quota '{quota_id}'). NO es que lo hayas agotado: simplemente no está "
                        "disponible en el tier gratuito para tu cuenta/región. "
                        "Prueba OTRO modelo (Diagnóstico) o activa facturación.",
                        response=resp,
                    )

                if es_diario:
                    _log("🚫 Límite DIARIO agotado — reintentar no ayudará hoy.")
                    raise requests.HTTPError(
                        f"429-DIARIO — Cuota DIARIA agotada (quota '{quota_id}', "
                        f"límite {quota_value}). {msg_api[:150]} "
                        "Soluciones: prueba otro modelo, espera 24h, o activa facturación.",
                        response=resp,
                    )

                if intento < len(esperas) - 1:
                    espera = esperas[intento] + round(random.uniform(0, 3), 1)
                    _log(f"⏳ Rate limit por minuto (429). Esperando {espera}s "
                         f"(backoff + jitter) antes de reintentar...")
                    time.sleep(espera)
                    continue
                raise requests.HTTPError(
                    f"429 — Rate limit tras varios intentos. {msg_api[:200]} "
                    "Espera 1-2 minutos y reintenta.",
                    response=resp,
                )

            if resp.status_code != 200:
                try:
                    err = resp.json().get("error", {}).get("message", resp.text[:200])
                except Exception:  # noqa: BLE001
                    err = resp.text[:200]
                _log(f"❌ [{version}] HTTP {resp.status_code}: {err}")
                resp.raise_for_status()

            _log(f"✅ Respuesta recibida en {dt:.1f}s. Leyendo contenido...")
            data = resp.json()
            try:
                cand = data["candidates"][0]
            except (KeyError, IndexError):
                fb = data.get("promptFeedback", {})
                _log(f"⚠️ Sin candidates. promptFeedback={json.dumps(fb)[:200]}")
                raise ValueError(
                    "Gemini no devolvió contenido (posible bloqueo de seguridad). "
                    f"{json.dumps(fb)[:200]}"
                )

            # Modelos 'thinking' devuelven varias parts: omitir las 'thought'.
            parts = cand.get("content", {}).get("parts", [])
            textos = [p["text"] for p in parts
                      if isinstance(p, dict) and "text" in p and not p.get("thought")]
            texto_final = "\n".join(textos).strip()

            finish = cand.get("finishReason", "")
            if finish and finish not in ("STOP", "MAX_TOKENS"):
                _log(f"⚠️ finishReason={finish}")
            if finish == "MAX_TOKENS":
                _log("⚠️ Respuesta truncada (MAX_TOKENS). El modelo 'pensó' demasiado. "
                     "Prueba un modelo más ligero (gemini-2.5-flash-lite) o divide el contenido.")

            if not texto_final:
                _log(f"⚠️ Respuesta sin texto utilizable. finishReason={finish}. "
                     f"Parts recibidas: {len(parts)}")
                raise ValueError(
                    f"Gemini respondió pero sin texto (finishReason={finish}). "
                    "Si es un modelo 'thinking', prueba gemini-2.5-flash-lite."
                )

            return texto_final

    raise requests.HTTPError(
        f"404 — El modelo '{modelo}' no existe en las APIs v1 ni v1beta. "
        "Prueba con 'gemini-2.5-flash' o 'gemini-2.5-flash-lite'."
    )


def _extraer_json(raw: str) -> dict:
    """Extrae un objeto JSON de la respuesta (quita fences, recorta {...})."""
    txt = raw.strip()
    if "```" in txt:
        for marker in ["```json", "```JSON", "```"]:
            if marker in txt:
                txt = txt.split(marker, 1)[1]
                if "```" in txt:
                    txt = txt.split("```")[0]
                txt = txt.strip()
                break
    try:
        return json.loads(txt)
    except json.JSONDecodeError:
        inicio = txt.find("{")
        fin = txt.rfind("}")
        if inicio != -1 and fin != -1 and fin > inicio:
            return json.loads(txt[inicio:fin + 1])
        raise


def aplicar_correccion_gemini(api_key: str, modelo: str,
                              datos_actuales: dict, peticion: str, log=None) -> dict:
    prompt = _PROMPT_CORRECCION.format(
        json_actual=json.dumps(datos_actuales, ensure_ascii=False, indent=2),
        peticion=peticion,
    )
    raw = _gemini_call(api_key, modelo, [prompt], log=log)
    return _extraer_json(raw)


def extraer_con_gemini_multi(api_key: str, modelo: str = MODELO_POR_DEFECTO,
                             texto: str = "", archivos: list | None = None,
                             log=None) -> dict:
    """Procesa texto + archivos (imágenes/PDF/Word) en una sola llamada."""
    def _log(msg: str):
        if log:
            log(msg)

    partes: list = [_PROMPT_GEMINI]
    texto_acumulado = texto.strip()

    for archivo in (archivos or []):
        datos = archivo.read()
        ext = archivo.name.rsplit(".", 1)[-1].lower()
        if ext in ("docx", "doc"):
            _log(f"📄 Extrayendo texto de {archivo.name}...")
            try:
                texto_doc = _texto_de_docx(datos)
                texto_acumulado += f"\n\n--- {archivo.name} ---\n{texto_doc}"
                _log(f"   → {len(texto_doc)} caracteres extraídos de {archivo.name}")
            except Exception as e:  # noqa: BLE001
                _log(f"   ⚠️ No se pudo leer {archivo.name}: {e}")
        elif ext == "pdf":
            _log(f"📎 Adjuntando PDF {archivo.name} ({len(datos)//1024} KB)...")
            partes.append({"mime_type": "application/pdf", "data": datos})
        elif ext in ("jpg", "jpeg", "png", "webp"):
            _log(f"🖼️ Adjuntando imagen {archivo.name} ({len(datos)//1024} KB)...")
            mime = archivo.type or f"image/{ext.replace('jpg', 'jpeg')}"
            partes.append({"mime_type": mime, "data": datos})

    if texto_acumulado:
        partes.append(f"\n\n--- TEXTO / DOCUMENTOS ---\n{texto_acumulado}")

    if len(partes) == 1:
        raise ValueError("No hay contenido para procesar.")

    raw = _gemini_call(api_key, modelo, partes, log=log)
    _log("🔍 Parseando preguntas del JSON devuelto...")
    vista = raw.strip().replace("\n", " ")
    _log(f"📄 Respuesta de Gemini (primeros 500 car.): {vista[:500]}")

    datos = _extraer_json(raw)
    n = len(datos.get("preguntas", []))
    _log(f"   → JSON parseado. Título: '{datos.get('titulo', '?')}' · "
         f"{n} pregunta(s) encontrada(s).")
    return datos


def gemini_a_tests(datos: dict) -> list[dict]:
    """Convierte la salida JSON de Gemini al formato canónico de la app."""
    titulo = datos.get("titulo", "Test importado con IA")
    preguntas: list[dict] = []
    for p in datos.get("preguntas", []):
        correcta = (p.get("correcta") or "").strip()
        if not correcta:
            continue
        incorrectas = [i.strip() for i in p.get("incorrectas", []) if i.strip()]
        opciones = [(correcta, True)] + [(inc, False) for inc in incorrectas]
        preguntas.append({
            "enunciado": (p.get("enunciado") or "").strip(),
            "opciones": opciones,
            "img_bytes": None,
            "img_nombre": None,
        })
    return [{"titulo": titulo, "preguntas": preguntas}]
