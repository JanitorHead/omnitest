"""Exportadores de los tests a Word, PDF, RemNote, Anki e imágenes,
más construir_resultado() que orquesta todos los formatos.
"""
import os
import tempfile
import zipfile
from collections.abc import Callable
from io import BytesIO

from .daypo import completar_imagenes_tests

ProgresoFn = Callable[[int, int, str], None]

import genanki
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF

from .quiz import generar_html_quiz
from .utils import generar_nombre_base, nombre_archivo, separar_opciones


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def _añadir_pregunta_word(doc: Document, num: int, pregunta: dict,
                          con_respuesta: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {pregunta['enunciado']}")
    run.bold = True
    run.font.size = Pt(11)

    if pregunta.get("img_bytes"):
        doc.add_picture(BytesIO(pregunta["img_bytes"]), width=Inches(3.5))

    for texto_op, es_correcta in pregunta["opciones"]:
        p_op = doc.add_paragraph()
        if con_respuesta and es_correcta:
            r = p_op.add_run(f"   - {texto_op}  (correcta)")
            r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x73, 0x28)
        else:
            p_op.add_run(f"   - {texto_op}")
    doc.add_paragraph()


def _doc_de_test(test: dict, con_respuesta: bool) -> bytes:
    doc = Document()
    doc.add_heading(test["titulo"], 1)
    for i, p in enumerate(test["preguntas"], 1):
        _añadir_pregunta_word(doc, i, p, con_respuesta=con_respuesta)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_word_combinado(tests: list[dict], con_respuesta: bool = True) -> bytes:
    if len(tests) == 1:
        return _doc_de_test(tests[0], con_respuesta)

    doc = Document()
    doc.add_heading("Recopilación de tests — Omnitest", 0)
    for test in tests:
        doc.add_heading(test["titulo"], 1)
        for i, p in enumerate(test["preguntas"], 1):
            _añadir_pregunta_word(doc, i, p, con_respuesta=con_respuesta)
        doc.add_page_break()
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_zip_word_individuales(tests: list[dict]) -> bytes:
    """ZIP con un .docx por test, en Con_Respuesta/ y Sin_Respuesta/."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for t in tests:
            nb = nombre_archivo(t["titulo"])
            zf.writestr(f"Con_Respuesta/{nb}.docx", _doc_de_test(t, True))
            zf.writestr(f"Sin_Respuesta/{nb}.docx", _doc_de_test(t, False))
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def generar_pdf(tests: list[dict], con_respuesta: bool = True) -> bytes:
    def safe(t: str) -> str:
        return t.encode("latin-1", errors="replace").decode("latin-1")

    MARGEN, W = 20, 170
    pdf = FPDF(format="A4")
    pdf.set_margins(MARGEN, MARGEN, MARGEN)
    pdf.set_auto_page_break(auto=True, margin=MARGEN)

    for test in tests:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.set_x(MARGEN)
        try:
            pdf.multi_cell(W, 10, safe(test["titulo"]), align="L")
        except Exception:
            pass
        pdf.ln(4)

        for i, p in enumerate(test["preguntas"], 1):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_x(MARGEN)
            try:
                pdf.multi_cell(W, 8, safe(f"{i}. {p['enunciado']}"), align="L")
            except Exception:
                pass

            if p.get("img_bytes"):
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
                    f.write(p["img_bytes"])
                    tmp = f.name
                try:
                    pdf.image(tmp, x=MARGEN, w=100)
                    pdf.ln(2)
                except Exception:
                    pass
                finally:
                    os.unlink(tmp)

            for texto, es_correcta in p["opciones"]:
                if not texto:
                    continue
                pdf.set_x(MARGEN + 4)
                if con_respuesta and es_correcta:
                    pdf.set_font("Helvetica", "B", 11)
                    contenido = safe(f">> {texto}")
                    ancho = W - 4
                else:
                    pdf.set_font("Helvetica", "", 11)
                    contenido = safe(texto)
                    ancho = W - 4
                try:
                    pdf.multi_cell(ancho, 7, contenido, align="L")
                except Exception:
                    pass
            pdf.ln(4)

    return bytes(pdf.output())


def generar_zip_pdf_individuales(tests: list[dict]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for t in tests:
            nb = nombre_archivo(t["titulo"])
            zf.writestr(f"Con_Respuesta/{nb}.pdf", generar_pdf([t], con_respuesta=True))
            zf.writestr(f"Sin_Respuesta/{nb}.pdf", generar_pdf([t], con_respuesta=False))
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# RemNote e imágenes
# ---------------------------------------------------------------------------

def generar_zip_remnote(tests: list[dict]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        lineas: list[str] = []
        for test in tests:
            lineas.append(test["titulo"])
            lineas.append("")
            for p in test["preguntas"]:
                correcta, incorrectas = separar_opciones(p["opciones"])
                if correcta is None:
                    continue
                if p.get("img_bytes") and p.get("img_nombre"):
                    nombre = p["img_nombre"]
                    zf.writestr(f"images/{nombre}", p["img_bytes"])
                    lineas.append(f"- **{p['enunciado']}")
                    lineas.append(f"**![](images/{nombre}) >>A)")
                else:
                    lineas.append(f"- **{p['enunciado']}** >>A)")
                lineas.append(f"    - {correcta}")
                for inc in incorrectas:
                    lineas.append(f"    - {inc}")
                lineas.append("")
            lineas.append("")
        zf.writestr("Banco_de_Preguntas_MCQ.md", "\n".join(lineas).encode("utf-8"))
    buf.seek(0)
    return buf.getvalue()


def generar_zip_imagenes(tests: list[dict]) -> bytes:
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for test in tests:
            carpeta = nombre_archivo(test["titulo"])
            n = 1
            for p in test["preguntas"]:
                if p.get("img_bytes"):
                    zf.writestr(f"{carpeta}/{carpeta}_{n:03d}.jpg", p["img_bytes"])
                    n += 1
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Anki
# ---------------------------------------------------------------------------

_ANKI_MODEL = genanki.Model(
    1607392324,
    "Omnitest MCQ Interactive",
    fields=[
        {"name": "Question"}, {"name": "Image"},
        {"name": "Option A"}, {"name": "Option B"}, {"name": "Option C"},
        {"name": "Option D"}, {"name": "Option E"}, {"name": "Correct Answer"},
    ],
    templates=[{
        "name": "MCQ Card",
        "qfmt": """\
<div class="card-container">
  <div class="question">{{Question}}</div>
  {{#Image}}<div class="image">{{Image}}</div>{{/Image}}
  <hr>
  <div class="options-container" id="options-box">
    {{#Option A}}<button class="option-btn" id="btn-A" onclick="selectOption('A')">{{Option A}}</button>{{/Option A}}
    {{#Option B}}<button class="option-btn" id="btn-B" onclick="selectOption('B')">{{Option B}}</button>{{/Option B}}
    {{#Option C}}<button class="option-btn" id="btn-C" onclick="selectOption('C')">{{Option C}}</button>{{/Option C}}
    {{#Option D}}<button class="option-btn" id="btn-D" onclick="selectOption('D')">{{Option D}}</button>{{/Option D}}
    {{#Option E}}<button class="option-btn" id="btn-E" onclick="selectOption('E')">{{Option E}}</button>{{/Option E}}
  </div>
</div>
<script>
function selectOption(letter) {
  var btns = document.getElementsByClassName("option-btn");
  for (var i = 0; i < btns.length; i++) btns[i].classList.remove("selected");
  document.getElementById("btn-" + letter).classList.add("selected");
  sessionStorage.setItem("ankiUserChoice", letter);
}
(function initOptions() {
  var box = document.getElementById("options-box");
  var btns = Array.from(box.children);
  var saved = sessionStorage.getItem("ankiOrder");
  if (saved) {
    saved.split(",").forEach(function(id) { var el=document.getElementById(id); if(el)box.appendChild(el); });
  } else {
    sessionStorage.removeItem("ankiUserChoice");
    for (var i=btns.length-1;i>0;i--){var j=Math.floor(Math.random()*(i+1));var t=btns[i];btns[i]=btns[j];btns[j]=t;}
    sessionStorage.setItem("ankiOrder", btns.map(function(b){return b.id;}).join(","));
    btns.forEach(function(b){box.appendChild(b);});
  }
})();
</script>""",
        "afmt": """\
{{FrontSide}}
<div id="correct-key" style="display:none;">{{Correct Answer}}</div>
<script>
var correct=document.getElementById("correct-key").innerText.trim().toUpperCase();
var chosen=sessionStorage.getItem("ankiUserChoice")||"";
var cBtn=document.getElementById("btn-"+correct);
if(cBtn)cBtn.classList.add("correct");
if(chosen&&chosen!==correct){var wBtn=document.getElementById("btn-"+chosen);if(wBtn)wBtn.classList.add("incorrect");}
sessionStorage.removeItem("ankiOrder");
</script>""",
    }],
    css="""\
.card{font-family:Arial,Helvetica,sans-serif;font-size:16px;text-align:left;color:#1a1a1a;background-color:#f0f0f0;}
.card-container{max-width:620px;margin:20px auto;padding:20px;background:#fff;border-radius:10px;box-shadow:0 4px 10px rgba(0,0,0,.12);}
.question{font-size:1.2em;font-weight:bold;margin-bottom:14px;color:#1a1a1a;}
.image{margin:12px 0;}.image img{max-width:100%;max-height:280px;border-radius:6px;}
.options-container{display:flex;flex-direction:column;gap:10px;}
.option-btn{text-align:left;padding:13px 16px;border:2px solid #d0d0d0;border-radius:8px;background:#fafafa;color:#1a1a1a;cursor:pointer;font-size:1em;transition:all .18s ease;}
.option-btn:hover{background:#efefef;border-color:#aaa;}
.option-btn.selected{border-color:#0056b3;background:#dceeff;color:#003a80;}
.option-btn.correct{border-color:#28a745!important;background:#d4edda!important;color:#155724!important;}
.option-btn.incorrect{border-color:#dc3545!important;background:#f8d7da!important;color:#721c24!important;}
.card.nightMode{color:#e8e8e8;background-color:#1e1e1e;}
.card.nightMode .card-container{background:#2a2a2a;}
.card.nightMode .question{color:#e8e8e8;}
.card.nightMode .option-btn{background:#333;border-color:#555;color:#e8e8e8;}
.card.nightMode .option-btn:hover{background:#3d3d3d;border-color:#777;}
.card.nightMode .option-btn.selected{border-color:#5aabff;background:#1a3a5c;color:#a8d4ff;}
.card.nightMode .option-btn.correct{border-color:#2ecc71!important;background:#1a3d2b!important;color:#7fdca0!important;}
.card.nightMode .option-btn.incorrect{border-color:#e74c3c!important;background:#3d1a1a!important;color:#f5a0a0!important;}
""",
)


def generar_apkg_anki(tests: list[dict]) -> bytes:
    deck = genanki.Deck(2059400110, "Omnitest")
    with tempfile.TemporaryDirectory() as tmpdir:
        media_files: list[str] = []
        for test in tests:
            for p in test["preguntas"]:
                correcta, incorrectas = separar_opciones(p["opciones"])
                if correcta is None:
                    continue
                campos = [correcta] + incorrectas
                fields_ops = [campos[i] if i < len(campos) else "" for i in range(5)]
                image_html = ""
                if p.get("img_bytes") and p.get("img_nombre"):
                    nombre = p["img_nombre"]
                    img_path = os.path.join(tmpdir, nombre)
                    with open(img_path, "wb") as f:
                        f.write(p["img_bytes"])
                    media_files.append(img_path)
                    image_html = f'<img src="{nombre}">'
                note = genanki.Note(
                    model=_ANKI_MODEL,
                    fields=[p["enunciado"], image_html, *fields_ops, "A"],
                    guid=genanki.guid_for(test["titulo"], p["enunciado"]),
                )
                deck.add_note(note)
        pkg = genanki.Package(deck)
        pkg.media_files = media_files
        apkg_path = os.path.join(tmpdir, "omnitest.apkg")
        pkg.write_to_file(apkg_path)
        with open(apkg_path, "rb") as f:
            return f.read()


# ---------------------------------------------------------------------------
# Orquestación
# ---------------------------------------------------------------------------

def construir_resultado(
    tests: list[dict],
    errores: list[str] | None = None,
    progreso: ProgresoFn | None = None,
) -> dict:
    """Genera todos los exportables y los empaqueta para session_state."""
    nb = generar_nombre_base(tests)
    es_multi = len(tests) > 1

    pasos: list[tuple[str, str, Callable[[], bytes]]] = []
    if any(t.get("id_test") for t in tests):
        pasos.append(("__imagenes__", "Completando imágenes…", lambda: b""))
    pasos.extend([
        ("word_combinado_con", "Generando Word (con respuestas)…", lambda: generar_word_combinado(tests, True)),
        ("word_combinado_sin", "Generando Word (sin respuestas)…", lambda: generar_word_combinado(tests, False)),
        ("pdf_combinado_con", "Generando PDF (con respuestas)…", lambda: generar_pdf(tests, True)),
        ("pdf_combinado_sin", "Generando PDF (sin respuestas)…", lambda: generar_pdf(tests, False)),
    ])
    if es_multi:
        pasos.extend([
            ("zip_word_individuales", "Generando Word por test…", lambda: generar_zip_word_individuales(tests)),
            ("zip_pdf_individuales", "Generando PDF por test…", lambda: generar_zip_pdf_individuales(tests)),
        ])
    pasos.extend([
        ("zip_imagenes", "Empaquetando imágenes…", lambda: generar_zip_imagenes(tests)),
        ("zip_remnote", "Generando RemNote…", lambda: generar_zip_remnote(tests)),
        ("apkg_anki", "Generando Anki…", lambda: generar_apkg_anki(tests)),
        ("html_quiz", "Generando quiz HTML…", lambda: generar_html_quiz(tests, nb)),
    ])

    total = len(pasos)
    datos: dict[str, bytes] = {}
    for i, (clave, etiqueta, fn) in enumerate(pasos, start=1):
        if progreso:
            progreso(i, total, etiqueta)
        if clave == "__imagenes__":
            completar_imagenes_tests(tests)
            continue
        datos[clave] = fn()

    return {
        "nombre_base": nb,
        "es_multi": es_multi,
        "word_combinado_con": datos["word_combinado_con"],
        "word_combinado_sin": datos["word_combinado_sin"],
        "zip_word_individuales": datos.get("zip_word_individuales", b""),
        "pdf_combinado_con": datos["pdf_combinado_con"],
        "pdf_combinado_sin": datos["pdf_combinado_sin"],
        "zip_pdf_individuales": datos.get("zip_pdf_individuales", b""),
        "zip_imagenes": datos["zip_imagenes"],
        "zip_remnote": datos["zip_remnote"],
        "apkg_anki": datos["apkg_anki"],
        "html_quiz": datos["html_quiz"],
        "tests_ok": len(tests),
        "errores": errores or [],
    }
