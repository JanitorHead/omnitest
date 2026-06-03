"""Capa unificada de proveedores de IA para extracción y corrección de tests."""
from __future__ import annotations

import json
from io import BytesIO

import requests
from docx import Document

from . import ai_import
from .api_config import PROVEEDORES, proveedor_activo, datos_proveedor, proveedores_configurados

PROMPT_EXTRACCION = ai_import._PROMPT_GEMINI  # noqa: SLF001
PROMPT_CORRECCION = ai_import._PROMPT_CORRECCION  # noqa: SLF001

_OPENAI_ENDPOINTS = {
    "groq": "https://api.groq.com/openai/v1/chat/completions",
    "cerebras": "https://api.cerebras.ai/v1/chat/completions",
    "mistral": "https://api.mistral.ai/v1/chat/completions",
}

_MODELS_ENDPOINTS = {
    "groq": "https://api.groq.com/openai/v1/models",
    "cerebras": "https://api.cerebras.ai/v1/models",
    "mistral": "https://api.mistral.ai/v1/models",
}

_EXCLUIR_MODELO = ("whisper", "embed", "tts", "guard", "moderation", "distil")


class ArchivoReleible:
    """Bytes de un upload re-leíbles (Streamlit solo permite leer una vez)."""

    __slots__ = ("name", "type", "_datos")

    def __init__(self, name: str, mime: str | None, datos: bytes):
        self.name = name
        self.type = mime
        self._datos = datos

    def read(self) -> bytes:
        return self._datos

    def seek(self, _pos: int = 0) -> None:
        pass


def _texto_de_docx(datos: bytes) -> str:
    doc = Document(BytesIO(datos))
    lineas = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            lineas.append("  |  ".join(c.text for c in fila.cells if c.text.strip()))
    return "\n".join(lineas)


def _preparar_contenido(texto: str, archivos: list | None, log=None) -> tuple[str, list[str]]:
    """Devuelve (texto_unificado, avisos) para proveedores solo-texto."""
    avisos: list[str] = []
    texto_acum = (texto or "").strip()

    for archivo in (archivos or []):
        datos = archivo.read()
        if not isinstance(datos, (bytes, bytearray)):
            datos = b""
        ext = archivo.name.rsplit(".", 1)[-1].lower()
        if ext in ("docx", "doc"):
            if log:
                log(f"📄 Extrayendo texto de {archivo.name}...")
            try:
                texto_doc = _texto_de_docx(datos)
                texto_acum += f"\n\n--- {archivo.name} ---\n{texto_doc}"
                if log:
                    log(f"   → {len(texto_doc)} caracteres de {archivo.name}")
            except Exception as e:  # noqa: BLE001
                avisos.append(f"No se pudo leer {archivo.name}: {e}")
        else:
            avisos.append(
                f"{archivo.name}: fotos/PDF requieren Google Gemini (multimodal). "
                "Convierte a texto o cambia de proveedor."
            )
    return texto_acum, avisos


def _texto_de_mensaje(message: dict) -> str:
    """Extrae texto de message.content o message.reasoning (p. ej. Cerebras gpt-oss)."""
    content = message.get("content")
    if isinstance(content, str) and content.strip():
        return content
    if isinstance(content, list):
        partes: list[str] = []
        for parte in content:
            if isinstance(parte, dict):
                t = parte.get("text") or parte.get("content")
                if t:
                    partes.append(str(t))
            elif isinstance(parte, str):
                partes.append(parte)
        unido = "".join(partes).strip()
        if unido:
            return unido
    reasoning = message.get("reasoning")
    if isinstance(reasoning, str) and reasoning.strip():
        return reasoning
    return ""


def _ping_openai_chat(pid: str, api_key: str, modelo: str) -> None:
    """Ping mínimo — éxito si HTTP 200 (sin exigir message.content)."""
    url = _OPENAI_ENDPOINTS[pid]
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    body = {
        "model": modelo,
        "messages": [{"role": "user", "content": "Responde solo: ok"}],
        "temperature": 0,
        "max_tokens": 32,
    }
    resp = requests.post(url, headers=headers, json=body, timeout=60)
    if resp.status_code == 429:
        raise requests.HTTPError(
            f"429 — Cuota o rate limit en {PROVEEDORES[pid]['nombre']}.",
            response=resp,
        )
    if resp.status_code != 200:
        try:
            err = resp.json().get("error", {}).get("message", resp.text[:200])
        except Exception:  # noqa: BLE001
            err = resp.text[:200]
        raise requests.HTTPError(
            f"HTTP {resp.status_code} ({PROVEEDORES[pid]['nombre']}): {err}",
            response=resp,
        )


def _openai_chat(
    pid: str,
    api_key: str,
    modelo: str,
    prompt: str,
    log=None,
    max_tokens: int = 8192,
    json_mode: bool | None = None,
) -> str:
    url = _OPENAI_ENDPOINTS[pid]
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    body = {
        "model": modelo,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
        "max_tokens": max_tokens,
    }
    use_json = json_mode if json_mode is not None else pid in ("groq", "mistral")
    if use_json:
        body["response_format"] = {"type": "json_object"}
    if log:
        log(f"🌐 Enviando a {PROVEEDORES[pid]['nombre']} ({modelo})...")
    resp = requests.post(url, headers=headers, json=body, timeout=120)
    if resp.status_code == 429:
        raise requests.HTTPError(
            f"429 — Cuota o rate limit en {PROVEEDORES[pid]['nombre']}. "
            "Espera un momento o prueba otro proveedor.",
            response=resp,
        )
    if resp.status_code != 200:
        try:
            err = resp.json().get("error", {}).get("message", resp.text[:200])
        except Exception:  # noqa: BLE001
            err = resp.text[:200]
        raise requests.HTTPError(
            f"HTTP {resp.status_code} ({PROVEEDORES[pid]['nombre']}): {err}",
            response=resp,
        )
    data = resp.json()
    message = data.get("choices", [{}])[0].get("message", {})
    texto = _texto_de_mensaje(message)
    if not texto:
        raise ValueError(f"Respuesta vacía de {PROVEEDORES[pid]['nombre']}")
    return texto


def _filtrar_modelos_chat(ids: list[str]) -> list[str]:
    out = [m for m in ids if not any(x in m.lower() for x in _EXCLUIR_MODELO)]
    return out or ids


def _modelo_para_prueba(pid: str, modelos: list[str]) -> str:
    from .model_router import MODELO_ROUTER

    preferido = MODELO_ROUTER.get(pid, "")
    if preferido in modelos:
        return preferido
    for pref in ("lite", "8b", "instant", "small", "mini"):
        for m in modelos:
            if pref in m.lower():
                return m
    return modelos[0] if modelos else preferido


def diagnosticar_proveedor(
    pid: str,
    api_key: str,
    modelo: str,
    modelos_api: list[str] | None = None,
) -> tuple[str, str]:
    """Ping mínimo. estado ∈ ok, cuota, no_existe, error."""
    if pid == "gemini":
        candidatos: list[str] = []
        for m in (modelo, *(modelos_api or [])):
            if m and m not in candidatos:
                candidatos.append(m)
        ultimo = ("no_existe", "modelo no encontrado")
        for m in candidatos:
            estado, det = ai_import.diagnosticar_modelo(api_key, m)
            if estado == "ok":
                return ("ok", "Funciona ✓" if m == modelo else f"Funciona ✓ ({m})")
            ultimo = (estado, det)
            if estado not in ("no_existe",):
                return ultimo
        return ultimo

    from .model_router import MODELOS_EXTRACCION

    candidatos = []
    for m in (modelo, *(modelos_api or []), *MODELOS_EXTRACCION.get(pid, [])):
        if m and m not in candidatos:
            candidatos.append(m)

    ultimo: requests.HTTPError | None = None
    for m in candidatos:
        try:
            _ping_openai_chat(pid, api_key, m)
            det = "Funciona ✓" if m == modelo else f"Funciona ✓ ({m})"
            return ("ok", det)
        except requests.HTTPError as e:
            ultimo = e
            if getattr(e.response, "status_code", None) == 404:
                continue
            if "429" in str(e):
                return ("cuota", "rate limit / cuota agotada")
            return ("error", str(e)[:120])
        except Exception as e:  # noqa: BLE001
            return ("error", str(e)[:120])

    if ultimo is not None:
        if "429" in str(ultimo):
            return ("cuota", "rate limit / cuota agotada")
        if getattr(ultimo.response, "status_code", None) == 404:
            return ("no_existe", "modelo no encontrado")
        return ("error", str(ultimo)[:120])
    return ("no_existe", "modelo no encontrado")


def listar_modelos(pid: str, api_key: str) -> list[str]:
    if pid == "gemini":
        return ai_import.listar_modelos_gemini(api_key)

    url = _MODELS_ENDPOINTS.get(pid)
    if not url:
        return list(PROVEEDORES[pid]["modelos_sugeridos"])

    headers = {"Authorization": f"Bearer {api_key}"}
    resp = requests.get(url, headers=headers, timeout=30)
    if resp.status_code == 401:
        raise requests.HTTPError("Clave API inválida", response=resp)
    resp.raise_for_status()
    ids = [m["id"] for m in resp.json().get("data", []) if m.get("id")]
    ids = _filtrar_modelos_chat(ids)
    return sorted(ids) if ids else list(PROVEEDORES[pid]["modelos_sugeridos"])


def probar_y_listar_proveedor(
    pid: str,
    api_key: str,
    config: dict | None = None,
) -> tuple[str, str, list[str]]:
    """Lista modelos vía API, prueba conectividad y opcionalmente persiste en config."""
    errores_lista: str | None = None
    try:
        modelos = listar_modelos(pid, api_key)
    except Exception as e:  # noqa: BLE001
        modelos = list(PROVEEDORES[pid]["modelos_sugeridos"])
        errores_lista = str(e)[:80]

    if config is not None:
        config["providers"][pid]["modelos_disponibles"] = modelos

    modelo_test = _modelo_para_prueba(pid, modelos)
    estado, det = diagnosticar_proveedor(pid, api_key, modelo_test, modelos_api=modelos)

    if modelos:
        det = f"{det} · {len(modelos)} modelo(s) recopilado(s)"
    if errores_lista:
        det = f"{det} (lista parcial: {errores_lista})"
    return estado, det, modelos


def validar_proveedores_en_config(config: dict) -> list[dict]:
    """Prueba y recopila modelos de cada proveedor con clave en config."""
    informes: list[dict] = []
    for pid in PROVEEDORES:
        key = (config["providers"][pid].get("api_key") or "").strip()
        if not key:
            continue
        estado, det, modelos = probar_y_listar_proveedor(pid, key, config)
        prov = config["providers"][pid]
        prov["habilitado"] = estado == "ok"
        informes.append({
            "pid": pid,
            "ok": estado == "ok",
            "det": det,
            "modelos": len(modelos),
        })
    return informes


def es_preview(modelo: str) -> bool:
    return ai_import.es_preview(modelo)


def extraer_preguntas(
    config: dict,
    texto: str = "",
    archivos: list | None = None,
    log=None,
    pid: str | None = None,
    modelo: str | None = None,
) -> dict:
    """Extrae preguntas con proveedor y modelo indicados (o activo por defecto)."""
    pid = pid or proveedor_activo(config)
    prov = datos_proveedor(config, pid)
    api_key = (prov.get("api_key") or "").strip()
    modelo = modelo or PROVEEDORES[pid]["modelo_defecto"]
    if not api_key:
        raise ValueError(f"No hay clave API para {PROVEEDORES[pid]['nombre']}.")

    if pid == "gemini":
        return ai_import.extraer_con_gemini_multi(
            api_key, modelo=modelo, texto=texto, archivos=archivos, log=log
        )

    texto_unificado, avisos = _preparar_contenido(texto, archivos, log=log)
    for av in avisos:
        if log:
            log(f"⚠️ {av}")
    if not texto_unificado.strip():
        raise ValueError(
            "No hay texto para procesar. "
            + (" ".join(avisos) if avisos else "Pega texto o sube un Word.")
        )

    prompt = PROMPT_EXTRACCION + f"\n\n--- TEXTO / DOCUMENTOS ---\n{texto_unificado}"
    raw = _openai_chat(pid, api_key, modelo, prompt, log=log)
    if log:
        log("🔍 Parseando JSON...")
    return ai_import._extraer_json(raw)  # noqa: SLF001


def aplicar_correccion(
    config: dict,
    datos_actuales: dict,
    peticion: str,
    log=None,
    pid: str | None = None,
    modelo: str | None = None,
) -> dict:
    pid = pid or proveedor_activo(config)
    prov = datos_proveedor(config, pid)
    api_key = (prov.get("api_key") or "").strip()
    modelo = modelo or PROVEEDORES[pid]["modelo_defecto"]
    prompt = PROMPT_CORRECCION.format(
        json_actual=json.dumps(datos_actuales, ensure_ascii=False, indent=2),
        peticion=peticion,
    )
    if pid == "gemini":
        raw = ai_import._gemini_call(api_key, modelo, [prompt], log=log)  # noqa: SLF001
    else:
        raw = _openai_chat(pid, api_key, modelo, prompt, log=log)
    return ai_import._extraer_json(raw)  # noqa: SLF001


def aplicar_correccion_con_fallback(
    config: dict,
    datos_actuales: dict,
    peticion: str,
    log=None,
    pid: str | None = None,
    modelo: str | None = None,
) -> tuple[dict, str, str]:
    """Corrige JSON probando el proveedor usado en extracción y luego el resto."""
    from .api_config import proveedor_listo, proveedores_configurados

    plan: list[tuple[str, str]] = []
    vistos: set[tuple[str, str]] = set()

    def _añadir(p: str, m: str) -> None:
        if proveedor_listo(config, p) and (p, m) not in vistos:
            plan.append((p, m))
            vistos.add((p, m))

    if pid and modelo:
        _añadir(pid, modelo)

    for p in proveedores_configurados(config):
        prov = config["providers"][p]
        modelos = prov.get("modelos_disponibles") or [PROVEEDORES[p]["modelo_defecto"]]
        for m in modelos:
            _añadir(p, m)
            break

    if not plan:
        raise ValueError("Configura al menos una API en ⚙ para corregir con IA.")

    ultimo: Exception | None = None
    for p, m in plan:
        try:
            datos = aplicar_correccion(
                config, datos_actuales, peticion, log=log, pid=p, modelo=m
            )
            return datos, p, m
        except Exception as e:  # noqa: BLE001
            ultimo = e
            if log:
                log(f"⚠️ {PROVEEDORES[p]['nombre']} falló — probando siguiente…")
            continue

    if ultimo:
        raise ultimo
    raise ValueError("No se pudo corregir con ningún proveedor configurado.")


def ia_a_tests(datos: dict) -> list[dict]:
    return ai_import.gemini_a_tests(datos)


def nombre_proveedor_activo(config: dict) -> str:
    pid = proveedor_activo(config)
    meta = PROVEEDORES.get(pid, {})
    return f"{meta.get('icono', '')} {meta.get('nombre', pid)}".strip()


def _normalizar_archivos(archivos: list | None) -> list | None:
    """Convierte uploads de Streamlit a objetos re-leíbles entre intentos."""
    if not archivos:
        return None
    normalizados: list[ArchivoReleible] = []
    for f in archivos:
        if isinstance(f, ArchivoReleible):
            normalizados.append(f)
            continue
        raw = f.read()
        if not isinstance(raw, (bytes, bytearray)):
            raw = b""
        normalizados.append(
            ArchivoReleible(f.name, getattr(f, "type", None), bytes(raw))
        )
    return normalizados


def extraer_con_fallback(
    config: dict,
    texto: str = "",
    archivos: list | None = None,
    log=None,
    progress=None,
) -> tuple[dict, str, bool, str, str]:
    """Router automático + extracción + fallback.

    Devuelve (datos, etiqueta_visible, hubo_fallback, pid_usado, modelo_usado).
    """
    from .model_router import analizar_metadatos_job, elegir_modelo, plan_fallback

    def _prog(p: int, msg: str = ""):
        if progress:
            progress(p)
        if log and msg:
            log(msg)

    archivos = _normalizar_archivos(archivos)
    if not proveedores_configurados(config):
        raise ValueError("Configura al menos una API en ⚙ para usar IA.")

    _prog(8, "Leyendo material…")
    meta = analizar_metadatos_job(texto, archivos)
    if meta["num_archivos"]:
        _prog(12, f"{meta['num_archivos']} archivo(s) · {meta['total_kb']} KB")
    if meta["texto_caracteres"]:
        _prog(15, f"Texto: {meta['texto_caracteres']} caracteres")

    _prog(18, "Analizando qué modelo conviene…")
    pid, modelo, reason = elegir_modelo(config, meta, log=log)
    _prog(25, f"Modelo elegido: {PROVEEDORES[pid]['nombre']} · {modelo}")
    if reason:
        _prog(28, reason)

    plan = plan_fallback(config, meta, pid, modelo)
    ultimo: Exception | None = None
    n_plan = max(len(plan), 1)
    proveedores_saltados: set[str] = set()

    for i, (p, m) in enumerate(plan):
        if p in proveedores_saltados:
            continue
        pct = 30 + int(55 * i / n_plan)
        _prog(pct, f"Extrayendo con {PROVEEDORES[p]['nombre']} ({m})…")
        try:
            datos = extraer_preguntas(config, texto, archivos, log=log, pid=p, modelo=m)
            nombre = PROVEEDORES[p]["nombre"]
            etiqueta = f"{nombre} · {m}"
            if i > 0:
                etiqueta += " (fallback)"
            _prog(88, "Parseando respuesta…")
            return datos, etiqueta, i > 0, p, m
        except Exception as e:  # noqa: BLE001
            ultimo = e
            err = str(e)
            fatal_cuota = (
                "429-DIARIO" in err
                or "429-CERO" in err
                or ("429" in err and p == "gemini" and not meta["multimodal_necesario"])
            )
            if fatal_cuota:
                proveedores_saltados.add(p)
                if log:
                    log(
                        f"⏭️ {PROVEEDORES[p]['nombre']} sin cuota — "
                        "usando otro proveedor…"
                    )
            elif log and i < len(plan) - 1:
                log(f"⚠️ {PROVEEDORES[p]['nombre']} falló — probando siguiente…")
            continue

    if ultimo:
        raise ultimo
    raise ValueError("No se pudo extraer con ningún proveedor configurado.")
